import json
from datetime import datetime, timedelta
from io import BytesIO
from unittest.mock import Mock, patch
from django.conf import settings
from django.contrib.auth import get_user_model
from django.core.exceptions import RequestDataTooBig
from django.db import connection
from django.test import RequestFactory, TestCase, override_settings
from django.test.utils import CaptureQueriesContext
from django.utils import timezone
from rest_framework.test import APIClient, force_authenticate
from assets.models import Camera, Server, System, Unit
from hechos.models import Hecho
from novedades.models import Novedad
from personnel.models import Person
from .models import FilmRecord, VideoAnalysisReport
from .services import IntegrityService
from .views import DashboardMapView, DashboardNovedadesView


_user_sequence = 0


def create_authenticated_user(*, username_prefix='tester', badge_prefix='88', role=Person.ROLE_ADMIN, with_person=True):
    global _user_sequence
    _user_sequence += 1
    username = f'{username_prefix}_{_user_sequence}'
    User = get_user_model()
    user = User.objects.create_user(username=username, password='password123!')
    user.is_staff = True
    user.is_superuser = True
    user.save(update_fields=['is_staff', 'is_superuser'])

    if with_person:
        badge_number = f'{badge_prefix}{_user_sequence:04d}'[-6:]
        Person.objects.create(
            first_name='Test',
            last_name='User',
            badge_number=badge_number,
            role=role,
            user=user,
        )
    return user


class DashboardLabelApiTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.user = create_authenticated_user(username_prefix='dashboard')
        self.client.force_authenticate(self.user)

    def test_personnel_dashboard_uses_humanized_role_labels(self):
        Person.objects.create(
            first_name='Ana',
            last_name='Control',
            badge_number='223456',
            role=Person.ROLE_COORDINADOR_COC,
        )
        Person.objects.create(
            first_name='Luis',
            last_name='Admin',
            badge_number='223457',
            role='ADMIN',
        )

        response = self.client.get('/api/dashboard/personnel/')

        self.assertEqual(response.status_code, 200)
        labels = {item['label'] for item in response.data['series']['distribution_primary']}
        self.assertIn(Person.ROLE_COORDINADOR_COC, labels)
        self.assertIn('ADMIN', labels)

    def test_novedades_dashboard_uses_humanized_status_labels(self):
        Novedad.objects.create(description='Camara sin servicio', status='OPEN')
        Novedad.objects.create(description='Servidor recuperado', status='CLOSED')

        response = self.client.get('/api/dashboard/novedades/')

        self.assertEqual(response.status_code, 200)
        labels = {item['label'] for item in response.data['series']['distribution_primary']}
        self.assertIn('OPEN', labels)
        self.assertIn('CLOSED', labels)

    def test_records_dashboard_exposes_finalizado_label(self):
        FilmRecord.objects.create(delivery_status='FINALIZADO')

        response = self.client.get('/api/dashboard/records/')

        self.assertEqual(response.status_code, 200)
        labels = {item['label'] for item in response.data['series']['distribution_primary']}
        self.assertIn('FINALIZADO', labels)

    def test_novedades_dashboard_trend_uses_selected_date_range(self):
        january = Novedad.objects.create(description='Incidente enero', status='OPEN')
        march = Novedad.objects.create(description='Incidente marzo', status='OPEN')

        Novedad.objects.filter(pk=january.pk).update(
            created_at=timezone.make_aware(datetime(2026, 1, 5, 10, 0, 0))
        )
        Novedad.objects.filter(pk=march.pk).update(
            created_at=timezone.make_aware(datetime(2026, 3, 5, 10, 0, 0))
        )

        response = self.client.get('/api/dashboard/novedades/?created_at__gte=2026-01-01&created_at__lte=2026-01-31')

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data['series']['trend'], [{'label': '2026-01-05', 'value': 1}])

    def test_records_dashboard_trend_uses_entry_date_period(self):
        FilmRecord.objects.create(delivery_status='ENTREGADO', entry_date=datetime(2026, 1, 7).date())
        FilmRecord.objects.create(delivery_status='FINALIZADO', entry_date=datetime(2026, 3, 2).date())

        response = self.client.get('/api/dashboard/records/?entry_date__gte=2026-01-01&entry_date__lte=2026-01-31')

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data['series']['trend'], [{'label': '2026-01-07', 'value': 1}])

    def test_personnel_dashboard_uses_selected_date_range(self):
        january = Person.objects.create(
            first_name='Ana',
            last_name='Enero',
            badge_number='323456',
            role=Person.ROLE_COORDINADOR_COC,
        )
        march = Person.objects.create(
            first_name='Luis',
            last_name='Marzo',
            badge_number='323457',
            role='ADMIN',
        )

        Person.objects.filter(pk=january.pk).update(
            created_at=timezone.make_aware(datetime(2026, 1, 10, 9, 0, 0))
        )
        Person.objects.filter(pk=march.pk).update(
            created_at=timezone.make_aware(datetime(2026, 3, 10, 9, 0, 0))
        )

        response = self.client.get('/api/dashboard/personnel/?created_at__gte=2026-01-01&created_at__lte=2026-01-31')

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data['totals']['records'], 1)
        self.assertEqual(response.data['series']['trend'], [{'label': '2026-01-10', 'value': 1}])


class DashboardQueryPerformanceTests(TestCase):
    def setUp(self):
        self.factory = RequestFactory()
        self.user = create_authenticated_user(username_prefix='queries')
        self.unit = Unit.objects.create(
            name='Ezeiza',
            code='EZE',
            airport='Ministro Pistarini',
            latitude='-34.8222',
            longitude='-58.5358',
            map_enabled=True,
        )
        system = System.objects.create(name='SYS-EZE', unit=self.unit)
        server = Server.objects.create(system=system, name='SRV-EZE', ip_address='10.0.0.10')
        camera = Camera.objects.create(server=server, name='CAM-EZE', ip_address='10.0.0.11')

        Novedad.objects.create(description='Novedad 1', status='OPEN', severity='HIGH', camera=camera)
        Novedad.objects.create(description='Novedad 2', status='CLOSED', severity='LOW', system=system)
        Hecho.objects.create(timestamp=timezone.now(), description='Hecho 1', camera=camera)
        FilmRecord.objects.create(generator_unit=self.unit, entry_date=timezone.localdate())
        Person.objects.create(
            first_name='Ana',
            last_name='Mapa',
            badge_number='555001',
            unit=self.unit,
            role=Person.ROLE_COORDINADOR_COC,
        )

    def test_novedades_dashboard_uses_a_small_fixed_number_of_queries(self):
        request = self.factory.get('/api/dashboard/novedades/')
        force_authenticate(request, user=self.user)
        view = DashboardNovedadesView.as_view()

        with CaptureQueriesContext(connection) as context:
            response = view(request)
            response.render()

        self.assertEqual(response.status_code, 200)
        self.assertLessEqual(len(context.captured_queries), 4)

    def test_map_dashboard_does_not_execute_queries_inside_the_unit_loop(self):
        request = self.factory.get('/api/dashboard/map/?scope=ba')
        force_authenticate(request, user=self.user)
        view = DashboardMapView.as_view()

        with CaptureQueriesContext(connection) as context:
            response = view(request)
            response.render()

        self.assertEqual(response.status_code, 200)
        self.assertLessEqual(len(context.captured_queries), 2)


class HechoFutureDateValidationApiTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.user = create_authenticated_user(username_prefix='hechos')
        self.client.force_authenticate(self.user)
        unit = Unit.objects.create(name='Unidad Hechos', code='UHT')
        system = System.objects.create(name='SYS-UHT', unit=unit)
        server = Server.objects.create(system=system, name='SRV-UHT', ip_address='10.0.1.10')
        self.camera = Camera.objects.create(server=server, name='CAM-UHT', ip_address='10.0.1.11')

    def test_rejects_future_timestamp(self):
        future_timestamp = (timezone.now() + timedelta(days=1)).isoformat()

        response = self.client.post(
            '/api/hechos/',
            {
                'timestamp': future_timestamp,
                'description': 'Hecho futuro',
                'camera_id': self.camera.id,
                'category': 'OPERATIVO',
            },
            format='json',
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('timestamp', response.data)

    def test_rejects_future_end_time(self):
        now = timezone.now()
        response = self.client.post(
            '/api/hechos/',
            {
                'timestamp': now.isoformat(),
                'end_time': (now + timedelta(hours=2)).isoformat(),
                'description': 'Hecho con cierre futuro',
                'camera_id': self.camera.id,
                'category': 'OPERATIVO',
            },
            format='json',
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('end_time', response.data)


class FilmRecordApiPeopleTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.user = create_authenticated_user(username_prefix='film-record')
        self.client.force_authenticate(self.user)
        self.records_url = '/api/film-records/'
        self.operator = Person.objects.create(
            first_name='Juan',
            last_name='Operador',
            badge_number='123456',
        )
        self.receiver = Person.objects.create(
            first_name='Maria',
            last_name='Recepcion',
            badge_number='123457',
        )
        self.unit = Unit.objects.create(name='Unidad Test', code='UTT')

    def _build_payload(self):
        return {
            "request_number": "SOL-2026-0001",
            "request_kind": "DENUNCIA",
            "request_type": "OFICIO",
            "judicial_case_number": "C-123/2026",
            "case_title": "DENUNCIA S/ HURTO DE EQUIPAJE",
            "judicial_office": "Fiscalia Nro. 1",
            "judicial_secretary": "",
            "judicial_holder": "",
            "criminal_problematic": "Hurto de equipaje",
            "incident_modality": "Sustraccion en cinta",
            "incident_place": "Hall de arribos",
            "incident_sector": "Cinta 3",
            "incident_time": "14:05",
            "generator_unit": self.unit.id,
            "operator": f"{self.operator.last_name}, {self.operator.first_name}",
            "received_by": self.receiver.id,
            "involved_people": [
                {
                    "role": "DENUNCIANTE",
                    "last_name": "Perez",
                    "first_name": "Ana",
                    "document_type": "DNI",
                    "document_number": "30111222",
                    "nationality": "Argentina",
                    "birth_date": "1990-06-10",
                },
                {
                    "role": "DETENIDO",
                    "last_name": "Gomez",
                    "first_name": "Luis",
                    "document_type": "DNI",
                    "document_number": "29888777",
                    "nationality": "Argentina",
                    "birth_date": "1988-01-05",
                },
            ],
        }

    def test_create_record_with_multiple_involved_people(self):
        payload = self._build_payload()
        response = self.client.post(self.records_url, payload, format='json')

        self.assertEqual(response.status_code, 201)
        self.assertIn('involved_people', response.data)
        self.assertEqual(len(response.data['involved_people']), 2)
        self.assertEqual(response.data['requester'], "Fiscalia Nro. 1")
        self.assertEqual(str(response.data['entry_date']), str(timezone.localdate()))

        record = FilmRecord.objects.get(pk=response.data['id'])
        self.assertEqual(record.involved_people.count(), 2)
        self.assertEqual(record.judicial_secretary, "N/C")
        self.assertEqual(record.judicial_holder, "N/C")

    def test_update_replaces_involved_people_without_orphans(self):
        create_response = self.client.post(self.records_url, self._build_payload(), format='json')
        self.assertEqual(create_response.status_code, 201)
        record_id = create_response.data['id']

        patch_payload = {
            "involved_people": [
                {
                    "role": "DENUNCIANTE",
                    "last_name": "Sosa",
                    "first_name": "Carla",
                    "document_type": "DNI",
                    "document_number": "33222444",
                    "nationality": "Argentina",
                    "birth_date": "1993-07-01",
                }
            ]
        }
        patch_response = self.client.patch(f'{self.records_url}{record_id}/', patch_payload, format='json')

        self.assertEqual(patch_response.status_code, 200)
        self.assertEqual(len(patch_response.data['involved_people']), 1)
        self.assertEqual(patch_response.data['involved_people'][0]['last_name'], 'Sosa')

        record = FilmRecord.objects.get(pk=record_id)
        self.assertEqual(record.involved_people.count(), 1)
        self.assertEqual(record.involved_people.first().last_name, 'Sosa')

    def test_missing_judicial_fields_are_normalized_to_nc(self):
        payload = self._build_payload()
        payload.pop('judicial_office', None)
        payload.pop('judicial_secretary', None)
        payload.pop('judicial_holder', None)
        payload['requester'] = ''

        response = self.client.post(self.records_url, payload, format='json')

        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data['judicial_office'], 'N/C')
        self.assertEqual(response.data['judicial_secretary'], 'N/C')
        self.assertEqual(response.data['judicial_holder'], 'N/C')
        self.assertEqual(response.data['requester'], 'N/C')


class FilmRecordFutureDateValidationApiTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.user = create_authenticated_user(username_prefix='film-record-future')
        self.client.force_authenticate(self.user)
        self.records_url = '/api/film-records/'
        self.unit = Unit.objects.create(name='Unidad Test', code='UTF')

    def _build_payload(self):
        return {
            "request_number": "SOL-2026-0001",
            "request_kind": "DENUNCIA",
            "request_type": "OFICIO",
            "judicial_case_number": "C-123/2026",
            "case_title": "DENUNCIA S/ HURTO DE EQUIPAJE",
            "judicial_office": "Fiscalia Nro. 1",
            "judicial_secretary": "",
            "judicial_holder": "",
            "criminal_problematic": "Hurto de equipaje",
            "incident_modality": "Sustraccion en cinta",
            "incident_place": "Hall de arribos",
            "incident_sector": "Cinta 3",
            "incident_time": "14:05",
            "generator_unit": self.unit.id,
            "operator": "Perez, Juan",
            "involved_people": [
                {
                    "role": "DENUNCIANTE",
                    "last_name": "Perez",
                    "first_name": "Ana",
                    "document_type": "DNI",
                    "document_number": "30111222",
                    "nationality": "Argentina",
                    "birth_date": "1990-06-10",
                }
            ],
        }

    def test_rejects_future_entry_date(self):
        payload = self._build_payload()
        payload['entry_date'] = str(timezone.localdate() + timedelta(days=1))

        response = self.client.post(self.records_url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('entry_date', response.data)

    def test_rejects_future_start_time(self):
        payload = self._build_payload()
        payload['start_time'] = (timezone.now() + timedelta(hours=3)).isoformat()

        response = self.client.post(self.records_url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('start_time', response.data)

    def test_rejects_future_birth_date_for_involved_people(self):
        payload = self._build_payload()
        payload['involved_people'][0]['birth_date'] = str(timezone.localdate() + timedelta(days=5))

        response = self.client.post(self.records_url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('involved_people', response.data)


class VideoAnalysisReportApiTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.user = create_authenticated_user(username_prefix='video-report')
        self.client.force_authenticate(self.user)
        self.url = '/api/video-analysis-report/'
        self.valid_report_data = {
            "report_date": "2026-03-09",
            "destinatarios": "Fiscalia Federal Nro. 2",
            "tipo_informe": "Informe de analisis de videos",
            "numero_informe": "0001EZE/2026",
            "grado": "OF. PRINCIPAL",
            "operador": "Perez, Juan",
            "lup": "506896",
            "sistema": "MILESTONE",
            "cantidad_observada": "2 personas",
            "sectores_analizados": "Hall de arribos, cinta 3, migraciones",
            "franja_horaria_analizada": "14:05 a 15:42",
            "tiempo_total_analisis": "1 hora 37 minutos",
            "sintesis": "No se observa manipulacion posterior del bulto",
            "hash_algorithms": ["sha256", "sha512"],
            "hash_program": "HashMyFiles",
            "medida_seguridad_interna": "Auditoria interna del COC",
            "vms_authenticity_mode": "vms_propio",
            "vms_authenticity_detail": "",
            "prevencion_sumaria": "003BAR/2026",
            "caratula": "DENUNCIA S/ PRESUNTO HURTO",
            "fiscalia": "Fiscalia Federal Nro. 2",
            "fiscal": "Dr. Carlos Sosa",
            "denunciante": "Perez, Ana",
            "involved_people_summary": "DENUNCIANTE: Perez, Ana (DNI 30111222) | DETENIDO: Gomez, Luis (DNI 29888777)",
            "involved_people": [
                {
                    "role": "DENUNCIANTE",
                    "full_name": "Perez, Ana",
                    "document_type": "DNI",
                    "document_number": "30111222",
                    "nationality": "Argentina",
                    "birth_date": "1990-06-10",
                    "age": 35,
                },
                {
                    "role": "DETENIDO",
                    "full_name": "Gomez, Luis",
                    "document_type": "DNI",
                    "document_number": "29888777",
                    "nationality": "Argentina",
                    "birth_date": "1988-01-05",
                    "age": 38,
                },
            ],
            "vuelo": "WJ 3045",
            "objeto_denunciado": "RIVER PLATE",
            "desarrollo": "Texto de desarrollo",
            "conclusion": "Texto de conclusion",
            "firma": "Coordinador CReV I DEL ESTE",
        }

    def _build_frame(self, index=0, content_base64="aGVsbG8="):
        return {
            "id_temp": f"f{index}",
            "file_name": f"frame{index}.jpg",
            "mime_type": "image/jpeg",
            "content_base64": content_base64,
            "frame_time": (timezone.localtime() - timedelta(minutes=5)).strftime('%Y-%m-%dT%H:%M'),
            "description": f"Frame {index}",
            "order": index
        }

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_generates_docx_with_new_payload(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        payload = {"report_data": self.valid_report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        service_mock.assert_called_once()

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_supports_legacy_flat_payload(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')

        response = self.client.post(self.url, self.valid_report_data, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_allows_generating_docx_without_signature(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {
            **self.valid_report_data,
            "firma": "",
        }

        response = self.client.post(self.url, {"report_data": report_data, "frames": []}, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_allows_generating_docx_with_base64_signature(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {
            **self.valid_report_data,
            "firma": "data:image/png;base64," + ("a" * 400),
        }

        response = self.client.post(self.url, {"report_data": report_data, "frames": []}, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_accepts_integrity_and_authenticity_fields(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {
            **self.valid_report_data,
            "hash_algorithms": ["sha3", "sha256", "sha512"],
            "hash_program": "HASH MY FILE",
            "medida_seguridad_interna": "Control de auditoria del COC",
            "vms_authenticity_mode": "hash_preventivo",
        }
        payload = {"report_data": report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()

    def test_rejects_invalid_hash_algorithm_in_report_data(self):
        report_data = {**self.valid_report_data, "hash_algorithms": ["sha256", "md5"]}
        payload = {"report_data": report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)
        self.assertIn('hash_algorithms', str(response.data))

    def test_rejects_future_report_date(self):
        report_data = {
            **self.valid_report_data,
            "report_date": str(timezone.localdate() + timedelta(days=1)),
        }

        response = self.client.post(self.url, {"report_data": report_data, "frames": []}, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)
        self.assertIn('report_date', str(response.data))

    def test_rejects_future_frame_time(self):
        frame = self._build_frame(index=0)
        frame["frame_time"] = (timezone.localtime() + timedelta(hours=2)).strftime('%Y-%m-%dT%H:%M')

        response = self.client.post(
            self.url,
            {"report_data": self.valid_report_data, "frames": [frame]},
            format='json',
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)
        self.assertIn('frame_time', str(response.data))
    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_maps_legacy_sintesis_fields_into_unified_sintesis(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {
            **self.valid_report_data,
            "sintesis": "",
            "sintesis_conclusion": "Sintesis legacy para conclusion",
        }

        response = self.client.post(self.url, {"report_data": report_data, "frames": []}, format='json')

        self.assertEqual(response.status_code, 200)
        call_payload = service_mock.call_args.args[0]
        self.assertEqual(
            call_payload.get("report_data", {}).get("sintesis"),
            "Sintesis legacy para conclusion"
        )

    def test_requires_hash_algorithm_other_when_other_is_selected(self):
        report_data = {
            **self.valid_report_data,
            "hash_algorithms": ["otro"],
            "hash_algorithm_other": "",
        }

        response = self.client.post(self.url, {"report_data": report_data, "frames": []}, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('hash_algorithm_other', str(response.data))

    def test_accepts_native_vms_metadata_without_requiring_detail(self):
        report_data = {
            **self.valid_report_data,
            "vms_authenticity_mode": "vms_propio",
            "vms_native_hash_algorithms": ["otro"],
            "vms_native_hash_algorithm_other": "",
        }

        response = self.client.post(self.url, {"report_data": report_data, "frames": []}, format='json')

        self.assertEqual(response.status_code, 200)

    def test_requires_vms_authenticity_detail_when_mode_is_otro(self):
        report_data = {
            **self.valid_report_data,
            "vms_authenticity_mode": "otro",
            "vms_authenticity_detail": "",
        }
        payload = {"report_data": report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)
        self.assertIn('vms_authenticity_detail', str(response.data))

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_allows_empty_fiscalia_and_fiscal(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {**self.valid_report_data, "fiscalia": "", "fiscal": ""}
        payload = {"report_data": report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_accepts_involved_people_in_report_payload(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {
            **self.valid_report_data,
            "involved_people_summary": "DENUNCIANTE: Perez Ana (DNI 30111222) | DETENIDO: Gomez Luis (DNI 29888777)",
            "involved_people": [
                {
                    "role": "DENUNCIANTE",
                    "full_name": "Perez, Ana",
                    "document_type": "DNI",
                    "document_number": "30111222",
                    "nationality": "Argentina",
                    "birth_date": "1990-06-10",
                    "age": 35,
                },
                {
                    "role": "DETENIDO",
                    "full_name": "Gomez, Luis",
                    "document_type": "DNI",
                    "document_number": "29888777",
                    "nationality": "Argentina",
                    "birth_date": "1988-01-05",
                    "age": 38,
                },
            ],
        }
        payload = {"report_data": report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_derives_destinatarios_from_fiscalia_when_missing(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {**self.valid_report_data, "fiscalia": "Fiscalia Nro. 02"}
        report_data.pop("destinatarios", None)
        payload = {"report_data": report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()
        call_payload = service_mock.call_args.args[0]
        self.assertEqual(
            call_payload.get("report_data", {}).get("destinatarios"),
            "Fiscalia Nro. 02"
        )

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_overrides_manual_destinatarios_with_fiscalia(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {
            **self.valid_report_data,
            "destinatarios": "Dato legacy manual",
            "fiscalia": "Fiscalia Nro. 03",
        }
        payload = {"report_data": report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()
        call_payload = service_mock.call_args.args[0]
        self.assertEqual(
            call_payload.get("report_data", {}).get("destinatarios"),
            "Fiscalia Nro. 03"
        )

    @patch('records.views.IntegrityService.generate_video_analysis_docx')
    def test_accepts_legacy_location_fields_without_error(self, service_mock):
        service_mock.return_value = (BytesIO(b'docx-data'), 'informe.docx')
        report_data = {
            **self.valid_report_data,
            "unidad_aeroportuaria": "Unidad Legacy",
            "asiento": "Asiento Legacy",
        }
        payload = {"report_data": report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        service_mock.assert_called_once()

    def test_rejects_invalid_frame_base64(self):
        payload = {
            "report_data": self.valid_report_data,
            "frames": [self._build_frame(index=0, content_base64="%%%invalid%%%")]
        }

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)

    def test_rejects_frame_without_brief_description(self):
        frame = self._build_frame(index=0)
        frame["description"] = ""
        payload = {
            "report_data": self.valid_report_data,
            "frames": [frame]
        }

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('description', str(response.data))

    def test_generates_docx_without_mock(self):
        payload = {"report_data": self.valid_report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        content_head = b''.join(response.streaming_content)[:2]
        self.assertEqual(content_head, b'PK')

    def test_rejects_more_than_max_frames(self):
        max_frames = int(getattr(settings, 'VIDEO_REPORT_MAX_FRAMES', 30))
        payload = {
            "report_data": self.valid_report_data,
            "frames": [self._build_frame(index=i) for i in range(max_frames + 1)]
        }

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)
        self.assertIn(str(max_frames), str(response.data))

    @patch('records.serializers.base64.b64decode')
    def test_rejects_frame_larger_than_max_frame_size(self, decode_mock):
        max_frame_size = int(getattr(settings, 'VIDEO_REPORT_MAX_FRAME_SIZE_BYTES', 8 * 1024 * 1024))
        decode_mock.return_value = b'a' * (max_frame_size + 1)

        payload = {
            "report_data": self.valid_report_data,
            "frames": [self._build_frame(index=0)]
        }

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)
        self.assertIn(f"{max_frame_size // (1024 * 1024)} MB", str(response.data))

    @patch('records.serializers.base64.b64decode')
    def test_rejects_total_frames_size_over_max_total(self, decode_mock):
        max_frames = int(getattr(settings, 'VIDEO_REPORT_MAX_FRAMES', 30))
        max_frame_size = int(getattr(settings, 'VIDEO_REPORT_MAX_FRAME_SIZE_BYTES', 8 * 1024 * 1024))
        max_total = int(getattr(settings, 'VIDEO_REPORT_MAX_TOTAL_BYTES', 80 * 1024 * 1024))

        chunk_size = min(max_frame_size, (max_total // max_frames) + 1)
        decode_mock.return_value = b'b' * chunk_size
        frame_count = (max_total // chunk_size) + 1

        payload = {
            "report_data": self.valid_report_data,
            "frames": [self._build_frame(index=i) for i in range(frame_count)]
        }

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)
        self.assertIn(f"{max_total // (1024 * 1024)} MB", str(response.data))

    @patch('records.views.VideoReportPayloadSerializer.is_valid', side_effect=RequestDataTooBig('too large'))
    def test_returns_413_when_request_data_is_too_large(self, _is_valid_mock):
        payload = {"report_data": self.valid_report_data, "frames": []}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 413)
        error_message = (response.data.get('error') or '').lower()
        self.assertIn('informe', error_message)
        self.assertIn('excede', error_message)


class VideoAnalysisReportStateApiTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.user = create_authenticated_user(username_prefix='report-state')
        self.client.force_authenticate(self.user)
        self.record = FilmRecord.objects.create()

    def test_save_report_draft_marks_report_as_borrador(self):
        response = self.client.post(
            f'/api/film-records/{self.record.id}/save_report_draft/',
            {
                'numero_informe': '0001EZE/2026',
                'report_date': '2026-03-10',
                'status': 'BORRADOR',
                'form_data': {'operador': 'Perez, Juan', 'grado': 'OF. PRINCIPAL'},
            },
            format='json',
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data['status'], 'BORRADOR')

        report = VideoAnalysisReport.objects.get(film_record=self.record)
        self.assertEqual(report.status, 'BORRADOR')
        self.assertEqual(report.numero_informe, '0001EZE/2026')

    def test_allows_updating_borrador_reports(self):
        report = VideoAnalysisReport.objects.create(
            film_record=self.record,
            numero_informe='0001EZE/2026',
            report_date=timezone.localdate(),
            status='BORRADOR',
            form_data={'operador': 'Perez, Juan'},
        )

        response = self.client.put(
            f'/api/video-analysis-reports/{report.id}/',
            {
                'film_record': self.record.id,
                'numero_informe': '0002EZE/2026',
                'report_date': '2026-03-11',
                'status': 'BORRADOR',
                'form_data': {'operador': 'Gomez, Ana'},
            },
            format='json',
        )

        self.assertEqual(response.status_code, 200)
        report.refresh_from_db()
        self.assertEqual(report.numero_informe, '0002EZE/2026')
        self.assertEqual(report.status, 'BORRADOR')
        self.assertEqual(report.form_data['operador'], 'Gomez, Ana')

    def test_rejects_updating_finalized_reports(self):
        report = VideoAnalysisReport.objects.create(
            film_record=self.record,
            numero_informe='0001EZE/2026',
            report_date=timezone.localdate(),
            status='FINALIZADO',
            form_data={'operador': 'Perez, Juan'},
        )

        response = self.client.put(
            f'/api/video-analysis-reports/{report.id}/',
            {
                'film_record': self.record.id,
                'numero_informe': '0003EZE/2026',
                'report_date': '2026-03-12',
                'status': 'BORRADOR',
                'form_data': {'operador': 'Cambio no permitido'},
            },
            format='json',
        )

        self.assertEqual(response.status_code, 403)
        report.refresh_from_db()
        self.assertEqual(report.status, 'FINALIZADO')
        self.assertEqual(report.numero_informe, '0001EZE/2026')

    def test_rejects_future_report_date_when_saving_draft(self):
        response = self.client.post(
            f'/api/film-records/{self.record.id}/save_report_draft/',
            {
                'numero_informe': '0001EZE/2026',
                'report_date': str(timezone.localdate() + timedelta(days=1)),
                'status': 'BORRADOR',
                'form_data': {'report_date': str(timezone.localdate() + timedelta(days=1))},
            },
            format='json',
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('report_date', str(response.data))


class IntegrityServiceVideoReportTemplateTests(TestCase):
    def test_resolves_informe_corto_template_first(self):
        template_path = IntegrityService._resolve_video_report_template_path()
        self.assertIn('INFORME CORTO 0010CREV-26 PS 003BAR-26.docx', str(template_path))

    def test_builds_verification_reference_from_hash_and_authenticity_data(self):
        reference = IntegrityService._build_verification_reference({
            "hash_program": "HashMyFiles",
            "hash_algorithms": ["sha256", "sha512"],
            "vms_authenticity_mode": "hash_preventivo",
            "vms_authenticity_detail": "",
        })
        self.assertIn('HashMyFiles', reference)
        self.assertIn('SHA-256', reference)
        self.assertIn('SHA-512', reference)
        self.assertIn('hash preventivo externo', reference)

    def test_builds_verification_reference_for_native_vms_without_operator_hash(self):
        reference = IntegrityService._build_verification_reference({
            "vms_authenticity_mode": "vms_propio",
            "vms_native_hash_algorithms": ["sha256"],
        })

        self.assertIn('propio sistema VMS', reference)
        self.assertNotIn('SHA-256', reference)
        self.assertNotIn('programa no consignado', reference)

    def test_builds_verification_reference_for_material_without_authentication(self):
        reference = IntegrityService._build_verification_reference({
            "vms_authenticity_mode": "sin_autenticacion",
            "motivo_sin_hash": "El sistema no ofrece autenticacion tecnica",
        })

        self.assertIn('sin autenticaci', reference)
        self.assertIn('Motivo: El sistema no ofrece autenticacion tecnica', reference)
        self.assertNotIn('Programa:', reference)

    def test_generate_video_analysis_docx_maps_informe_corto_layout(self):
        payload = {
            "report_data": {
                "report_date": "2026-03-05",
                "numero_informe": "0010CREV/2026",
                "fiscalia": "UFI 1 DESCENTRALIZADA",
                "fiscal": "Dra. Maria Bello",
                "prevencion_sumaria": "003BAR/2026",
                "caratula": "AMBAR JUSTBICHER S/ DENUNCIA",
                "denunciante": "AMBAR JUSTBICHER",
                "aeropuerto": "Terminal B",
                "fecha_hecho": "05/03/2026",
                "sintesis": "No se advierte nueva manipulacion del equipaje denunciado",
                "hash_program": "HashMyFiles",
                "hash_algorithms": ["sha256", "sha512"],
                "vms_authenticity_mode": "hash_preventivo",
                "material_filmico": "MATERIAL FILMICO TEST",
                "desarrollo": "DESARROLLO TEST",
                "conclusion": "CONCLUSION TEST",
            },
            "frames": [],
        }

        out, filename = IntegrityService.generate_video_analysis_docx(payload)
        self.assertEqual(filename, "INFORME CORTO 0010CREV-26 PS 003BAR-26.docx")

        from docx import Document

        document = Document(out)
        paragraphs_text = "\n".join((paragraph.text or "") for paragraph in document.paragraphs)
        self.assertIn("Referencia de verificación:", paragraphs_text)
        self.assertIn("MATERIAL FILMICO TEST", paragraphs_text)
        self.assertIn("DESARROLLO TEST", paragraphs_text)
        self.assertIn("CONCLUSION TEST", paragraphs_text)

        self.assertEqual(len(document.tables), 2)
        self.assertEqual(len(document.tables[0].rows), 6)
        self.assertEqual(len(document.tables[1].rows), 2)

        table_rows = [row.cells[0].text.strip() for row in document.tables[0].rows]
        self.assertIn("PREVENCION SUMARIA: 003BAR/2026", table_rows)
        self.assertIn("CARATULA: “AMBAR JUSTBICHER S/ DENUNCIA”", table_rows)
        self.assertIn("DAMNIFICADO: AMBAR JUSTBICHER. -", table_rows)
        self.assertIn("AEROPUERTO: TERMINAL B. -", table_rows)
        fiscal_rows = [row for row in table_rows if row.startswith("FISCALIA:")]
        self.assertTrue(fiscal_rows)


class SerializerLimitConsistencyTests(TestCase):
    """Verify that serializer fallback defaults match settings.py values."""

    def test_max_frames_default_matches_settings(self):
        from .serializers import _video_report_max_frames
        expected = int(getattr(settings, 'VIDEO_REPORT_MAX_FRAMES', 30))
        self.assertEqual(_video_report_max_frames(), expected)

    def test_max_frame_size_default_matches_settings(self):
        from .serializers import _video_report_max_frame_size_bytes
        expected = int(getattr(settings, 'VIDEO_REPORT_MAX_FRAME_SIZE_BYTES', 8 * 1024 * 1024))
        self.assertEqual(_video_report_max_frame_size_bytes(), expected)

    def test_max_total_bytes_default_matches_settings(self):
        from .serializers import _video_report_max_total_bytes
        expected = int(getattr(settings, 'VIDEO_REPORT_MAX_TOTAL_BYTES', 80 * 1024 * 1024))
        self.assertEqual(_video_report_max_total_bytes(), expected)


class VideoAnalysisImproveTextApiTests(TestCase):
    def setUp(self):
        self.client = APIClient()
        self.user = create_authenticated_user(username_prefix='improve-text')
        self.client.force_authenticate(self.user)
        self.url = '/api/video-analysis-improve-text/'

    @patch('records.views.IntegrityService.improve_report_text_with_ai')
    def test_improves_text_fields(self, improve_mock):
        improve_mock.return_value = {
            "material_filmico": "Material filmico mejorado",
            "desarrollo": "Desarrollo mejorado",
            "conclusion": "Conclusion mejorada",
        }
        payload = {
            "material_filmico": "Texto original de material filmico",
            "desarrollo": "Texto original de desarrollo",
            "conclusion": "Texto original de conclusion",
            "material_context": {
                "sistema": "MILESTONE",
                "aeropuerto": "Aeropuerto Internacional Mtro. Pistarini",
                "cantidad_observada": "2 personas",
                "sectores_analizados": "Hall de arribos, cinta 3",
                "franja_horaria_analizada": "14:05 a 15:42",
                "tiempo_total_analisis": "1 hora 37 minutos",
                "sintesis_conclusion": "No se observa manipulaciÃƒÂ³n posterior del bulto",
                "hash_algorithms": ["sha256", "sha512"],
                "vms_authenticity_mode": "vms_propio",
            },
        }

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data.get('material_filmico'), "Material filmico mejorado")
        self.assertEqual(response.data.get('desarrollo'), "Desarrollo mejorado")
        self.assertEqual(response.data.get('conclusion'), "Conclusion mejorada")
        improve_mock.assert_called_once_with(
            "Texto original de material filmico",
            "Texto original de desarrollo",
            "Texto original de conclusion",
            custom_api_key="",
            material_context={
                "sistema": "MILESTONE",
                "aeropuerto": "Aeropuerto Internacional Mtro. Pistarini",
                "cantidad_observada": "2 personas",
                "sectores_analizados": "Hall de arribos, cinta 3",
                "franja_horaria_analizada": "14:05 a 15:42",
                "tiempo_total_analisis": "1 hora 37 minutos",
                "sintesis": "No se observa manipulaciÃƒÂ³n posterior del bulto",
                "sintesis_conclusion": "No se observa manipulaciÃƒÂ³n posterior del bulto",
                "hash_algorithms": ["sha256", "sha512"],
                "vms_authenticity_mode": "vms_propio",
            },
            mode="full",
            preferred_provider="",
        )

    @patch('records.views.IntegrityService.improve_report_text_with_ai')
    def test_accepts_context_only_payload(self, improve_mock):
        improve_mock.return_value = {
            "material_filmico": "Material filmico mejorado",
            "desarrollo": "Desarrollo mejorado",
            "conclusion": "Conclusion mejorada",
        }

        response = self.client.post(
            self.url,
            {
                "material_filmico": "",
                "desarrollo": "",
                "conclusion": "",
                "material_context": {
                    "sectores_analizados": "Hall de arribos, cinta 3",
                    "franja_horaria_analizada": "14:05 a 15:42",
                    "tiempo_total_analisis": "1 hora 37 minutos",
                    "sintesis_conclusion": "No se observa manipulaciÃƒÂ³n posterior del bulto",
                }
            },
            format='json'
        )

        self.assertEqual(response.status_code, 200)
        improve_mock.assert_called_once()

    def test_rejects_empty_payload(self):
        response = self.client.post(
            self.url,
            {"material_filmico": "", "desarrollo": "", "conclusion": ""},
            format='json'
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)

    def test_rejects_context_when_vms_mode_otro_without_detail(self):
        response = self.client.post(
            self.url,
            {
                "material_filmico": "texto",
                "material_context": {
                    "vms_authenticity_mode": "otro",
                    "vms_authenticity_detail": "",
                },
            },
            format='json'
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('errors', response.data)
        self.assertIn('vms_authenticity_detail', str(response.data))

    @patch('records.views.IntegrityService.improve_report_text_with_ai', side_effect=RuntimeError('No se configuro AI_TEXT_API_KEY.'))
    def test_returns_503_when_ai_is_not_configured(self, _improve_mock):
        payload = {"desarrollo": "algo", "conclusion": ""}

        response = self.client.post(self.url, payload, format='json')

        self.assertEqual(response.status_code, 503)
        self.assertEqual(response.data.get('error'), 'No se configuro AI_TEXT_API_KEY.')


class IntegrityServiceAiRequestTests(TestCase):
    @staticmethod
    def _mock_ai_success_response():
        response_mock = Mock()
        response_mock.status_code = 200
        response_mock.json.return_value = {
            "choices": [
                {"message": {"content": '{"material_filmico":"Material filmico mejorado","desarrollo":"Desarrollo mejorado","conclusion":"Conclusion mejorada"}'}}
            ]
        }
        return response_mock

    @staticmethod
    def _mock_ai_error_response(status_code, message):
        response_mock = Mock()
        response_mock.status_code = status_code
        response_mock.json.return_value = {"error": {"message": message}}
        response_mock.text = message
        return response_mock

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='gemini,openrouter,groq',
        AI_TEXT_PROVIDER_SELECTION='round_robin',
        GEMINI_API_KEY='gemini-key',
        AI_TEXT_GEMINI_API_URL='https://generativelanguage.googleapis.com/v1beta/openai/chat/completions',
        AI_TEXT_GEMINI_MODEL='gemini-2.5-flash',
        OPEN_ROUTER_API_KEY='openrouter-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
        GROQ_API_KEY='groq-key',
        AI_TEXT_GROQ_API_URL='https://api.groq.com/openai/v1/chat/completions',
        AI_TEXT_GROQ_MODEL='llama-3.3-70b-versatile',
    )
    def test_round_robin_rotates_start_provider(self):
        IntegrityService._provider_rotation_index = 0

        first = IntegrityService._get_ai_provider_chain()
        second = IntegrityService._get_ai_provider_chain()
        third = IntegrityService._get_ai_provider_chain()

        self.assertEqual([item['name'] for item in first], ['gemini', 'openrouter', 'groq'])
        self.assertEqual([item['name'] for item in second], ['openrouter', 'groq', 'gemini'])
        self.assertEqual([item['name'] for item in third], ['groq', 'gemini', 'openrouter'])

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='openrouter',
        AI_TEXT_PROVIDER_SELECTION='ordered',
        AI_TEXT_FALLBACK_MODE='quota_only',
        OPEN_ROUTER_API_KEY='test-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
        AI_TEXT_TIMEOUT_SECONDS=33,
    )
    @patch('records.services.requests.post')
    def test_sends_only_narrative_fields_to_ai(self, post_mock):
        post_mock.return_value = self._mock_ai_success_response()

        result = IntegrityService.improve_report_text_with_ai(
            "Texto original de material filmico",
            "Texto original de desarrollo",
            "Texto original de conclusion",
        )

        self.assertEqual(result.get('material_filmico'), 'Material filmico mejorado')
        self.assertEqual(result.get('desarrollo'), 'Desarrollo mejorado')
        self.assertEqual(result.get('conclusion'), 'Conclusion mejorada')
        post_mock.assert_called_once()

        _, kwargs = post_mock.call_args
        self.assertEqual(kwargs.get('timeout'), 33)
        self.assertEqual(kwargs.get('headers', {}).get('Authorization'), 'Bearer test-key')

        request_payload = kwargs.get('json') or {}
        self.assertEqual(
            set(request_payload.keys()),
            {'model', 'temperature', 'messages', 'response_format'}
        )
        self.assertEqual(request_payload.get('model'), 'gpt-4o-mini')
        self.assertEqual(request_payload.get('response_format'), {'type': 'json_object'})

        messages = request_payload.get('messages') or []
        self.assertEqual(len(messages), 2)
        user_content = json.loads(messages[1].get('content') or '{}')
        self.assertEqual(
            set(user_content.keys()),
            {
                'instrucciones',
                'formato_estricto',
                'material_filmico_original',
                'desarrollo_original',
                'conclusion_original',
                'material_context',
            }
        )
        self.assertEqual(user_content.get('material_filmico_original'), 'Texto original de material filmico')
        self.assertEqual(user_content.get('desarrollo_original'), 'Texto original de desarrollo')
        self.assertEqual(user_content.get('conclusion_original'), 'Texto original de conclusion')
        self.assertEqual(user_content.get('material_context'), {})

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='openrouter',
        AI_TEXT_PROVIDER_SELECTION='ordered',
        AI_TEXT_FALLBACK_MODE='quota_only',
        OPEN_ROUTER_API_KEY='test-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
        AI_TEXT_TIMEOUT_SECONDS=33,
    )
    @patch('records.services.requests.post')
    def test_includes_material_context_in_ai_prompt_when_provided(self, post_mock):
        post_mock.return_value = self._mock_ai_success_response()

        context = {
            "sistema": "MILESTONE",
            "aeropuerto": "Aeropuerto Internacional Mtro. Pistarini",
            "cantidad_observada": "2 personas",
            "sectores_analizados": "Hall de arribos, cinta 3",
            "franja_horaria_analizada": "14:05 a 15:42",
            "tiempo_total_analisis": "1 hora 37 minutos",
            "sintesis_conclusion": "No se observa manipulaciÃƒÂ³n posterior del bulto",
            "hash_algorithms": ["sha256", "sha512"],
            "hash_program": "HASH MY FILE",
            "medida_seguridad_interna": "Auditoria interna",
            "vms_authenticity_mode": "vms_propio",
        }
        IntegrityService.improve_report_text_with_ai(
            "Texto material",
            "Texto desarrollo",
            "Texto conclusion",
            material_context=context,
        )

        _, kwargs = post_mock.call_args
        request_payload = kwargs.get('json') or {}
        messages = request_payload.get('messages') or []
        user_content = json.loads(messages[1].get('content') or '{}')
        self.assertEqual(user_content.get('material_context'), context)

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='gemini,openrouter,groq',
        AI_TEXT_PROVIDER_SELECTION='ordered',
        AI_TEXT_FALLBACK_MODE='quota_only',
        GEMINI_API_KEY='gemini-key',
        AI_TEXT_GEMINI_API_URL='https://generativelanguage.googleapis.com/v1beta/openai/chat/completions',
        AI_TEXT_GEMINI_MODEL='gemini-2.5-flash',
        OPEN_ROUTER_API_KEY='openrouter-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
        GROQ_API_KEY='groq-key',
        AI_TEXT_GROQ_API_URL='https://api.groq.com/openai/v1/chat/completions',
        AI_TEXT_GROQ_MODEL='llama-3.3-70b-versatile',
        AI_TEXT_TIMEOUT_SECONDS=20,
    )
    @patch('records.services.requests.post')
    def test_fallbacks_from_gemini_to_openrouter_on_quota(self, post_mock):
        post_mock.side_effect = [
            self._mock_ai_error_response(429, 'RESOURCE_EXHAUSTED: resource has been exhausted'),
            self._mock_ai_success_response(),
        ]

        result = IntegrityService.improve_report_text_with_ai(
            "texto material",
            "texto desarrollo",
            "texto conclusion"
        )

        self.assertEqual(result.get('material_filmico'), 'Material filmico mejorado')
        self.assertEqual(result.get('desarrollo'), 'Desarrollo mejorado')
        self.assertEqual(result.get('conclusion'), 'Conclusion mejorada')
        self.assertEqual(post_mock.call_count, 2)
        self.assertIn('generativelanguage.googleapis.com', post_mock.call_args_list[0].kwargs.get('url', ''))
        self.assertIn('openrouter.ai', post_mock.call_args_list[1].kwargs.get('url', ''))

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='gemini,openrouter,groq',
        AI_TEXT_PROVIDER_SELECTION='ordered',
        AI_TEXT_FALLBACK_MODE='quota_only',
        GEMINI_API_KEY='gemini-key',
        AI_TEXT_GEMINI_API_URL='https://generativelanguage.googleapis.com/v1beta/openai/chat/completions',
        AI_TEXT_GEMINI_MODEL='gemini-2.5-flash',
        OPEN_ROUTER_API_KEY='openrouter-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
        AI_TEXT_TIMEOUT_SECONDS=20,
    )
    @patch('records.services.requests.post')
    def test_does_not_fallback_on_non_quota_error(self, post_mock):
        post_mock.return_value = self._mock_ai_error_response(400, 'Bad request: invalid response_format')

        with self.assertRaises(RuntimeError) as exc:
            IntegrityService.improve_report_text_with_ai(
                "texto material",
                "texto desarrollo",
                "texto conclusion"
            )

        self.assertIn('La API de IA devolvio un error (400).', str(exc.exception))
        self.assertEqual(post_mock.call_count, 1)

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='gemini,openrouter,groq',
        AI_TEXT_PROVIDER_SELECTION='ordered',
        AI_TEXT_FALLBACK_MODE='quota_only',
        GEMINI_API_KEY='gemini-key',
        AI_TEXT_GEMINI_API_URL='https://generativelanguage.googleapis.com/v1beta/openai/chat/completions',
        AI_TEXT_GEMINI_MODEL='gemini-2.5-flash',
        OPEN_ROUTER_API_KEY='openrouter-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
        GROQ_API_KEY='groq-key',
        AI_TEXT_GROQ_API_URL='https://api.groq.com/openai/v1/chat/completions',
        AI_TEXT_GROQ_MODEL='llama-3.3-70b-versatile',
        AI_TEXT_TIMEOUT_SECONDS=20,
    )
    @patch('records.services.requests.post')
    def test_returns_controlled_error_when_all_providers_hit_quota(self, post_mock):
        post_mock.side_effect = [
            self._mock_ai_error_response(429, 'RESOURCE_EXHAUSTED'),
            self._mock_ai_error_response(402, 'insufficient credits'),
            self._mock_ai_error_response(429, 'rate limit exceeded'),
        ]

        with self.assertRaises(RuntimeError) as exc:
            IntegrityService.improve_report_text_with_ai(
                "texto material",
                "texto desarrollo",
                "texto conclusion"
            )

        self.assertEqual(
            str(exc.exception),
            'Todos los proveedores de IA alcanzaron su limite de cuota/tokens.'
        )
        self.assertEqual(post_mock.call_count, 3)

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='openrouter',
        AI_TEXT_PROVIDER_SELECTION='ordered',
        AI_TEXT_FALLBACK_MODE='quota_only',
        OPEN_ROUTER_API_KEY='test-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
    )
    @patch('records.services.requests.post')
    def test_mode_material_filmico_sends_system_data_context(self, post_mock):
        """mode='material_filmico' debe enviar datos tÃƒÂ©cnicos (sistema, hash, VMS)."""
        response_mock = Mock()
        response_mock.status_code = 200
        response_mock.json.return_value = {
            "choices": [{"message": {"content": '{"material_filmico":"Material generado"}'}}]
        }
        post_mock.return_value = response_mock

        context = {
            "sistema": "MILESTONE",
            "aeropuerto": "Aeroparque",
            "hash_algorithms": ["sha256"],
            "hash_program": "HashMyFiles",
            "vms_authenticity_mode": "vms_propio",
            "frames_summary": "Fotograma 1: fecha y hora 12/03/2026 15:00 - equipaje sobre la cinta",
        }
        result = IntegrityService.improve_report_text_with_ai(
            "texto original", "", "", material_context=context, mode='material_filmico'
        )

        self.assertEqual(result.get('material_filmico'), 'Material generado')
        # modo especÃƒÂ­fico: desarrollo y conclusion no cambian
        self.assertEqual(result.get('desarrollo'), '')
        self.assertEqual(result.get('conclusion'), '')

        _, kwargs = post_mock.call_args
        messages = kwargs.get('json', {}).get('messages', [])
        user_content = json.loads(messages[1].get('content') or '{}')
        self.assertIn('datos_del_sistema', user_content)
        self.assertEqual(user_content['datos_del_sistema']['sistema_cctv'], 'MILESTONE')
        self.assertEqual(user_content['datos_del_sistema']['algoritmos_hash'], 'SHA-256')
        self.assertEqual(user_content['datos_del_sistema']['medida_autenticidad'], 'VMS propio del sistema de grabacion')
        self.assertEqual(
            user_content['datos_del_sistema']['referencias_de_fotogramas'],
            'Fotograma 1: fecha y hora 12/03/2026 15:00 - equipaje sobre la cinta',
        )
        # Solo debe pedir material_filmico en el formato
        self.assertEqual(set(user_content['formato_estricto'].keys()), {'material_filmico'})

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='openrouter',
        AI_TEXT_PROVIDER_SELECTION='ordered',
        AI_TEXT_FALLBACK_MODE='quota_only',
        OPEN_ROUTER_API_KEY='test-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
    )
    @patch('records.services.requests.post')
    def test_mode_desarrollo_sends_analysis_context(self, post_mock):
        """mode='desarrollo' debe enviar datos analÃƒÂ­ticos (sectores, franja, tiempo, cantidad)."""
        response_mock = Mock()
        response_mock.status_code = 200
        response_mock.json.return_value = {
            "choices": [{"message": {"content": '{"desarrollo":"Desarrollo generado"}'}}]
        }
        post_mock.return_value = response_mock

        context = {
            "sectores_analizados": "Hall de arribos",
            "franja_horaria_analizada": "14:00 a 15:30",
            "tiempo_total_analisis": "1 hora 30 minutos",
            "cantidad_observada": "3 pasajeros",
            "frames_summary": "Fotograma 2: fecha y hora 12/03/2026 02:14 - el sujeto toma una mochila",
        }
        result = IntegrityService.improve_report_text_with_ai(
            "", "texto desarrollo", "", material_context=context, mode='desarrollo'
        )

        self.assertEqual(result.get('desarrollo'), 'Desarrollo generado')
        self.assertEqual(result.get('material_filmico'), '')
        self.assertEqual(result.get('conclusion'), '')

        _, kwargs = post_mock.call_args
        messages = kwargs.get('json', {}).get('messages', [])
        user_content = json.loads(messages[1].get('content') or '{}')
        self.assertIn('datos_del_analisis', user_content)
        self.assertEqual(user_content['datos_del_analisis']['sectores_analizados'], 'Hall de arribos')
        self.assertEqual(
            user_content['datos_del_analisis']['referencias_de_fotogramas'],
            'Fotograma 2: fecha y hora 12/03/2026 02:14 - el sujeto toma una mochila',
        )
        self.assertEqual(set(user_content['formato_estricto'].keys()), {'desarrollo'})

    @override_settings(
        AI_TEXT_PROVIDER_ORDER='openrouter',
        AI_TEXT_PROVIDER_SELECTION='ordered',
        AI_TEXT_FALLBACK_MODE='quota_only',
        OPEN_ROUTER_API_KEY='test-key',
        AI_TEXT_OPENROUTER_API_URL='https://openrouter.ai/api/v1/chat/completions',
        AI_TEXT_OPENROUTER_MODEL='gpt-4o-mini',
    )
    @patch('records.services.requests.post')
    def test_mode_conclusion_sends_synthesis_context(self, post_mock):
        """mode='conclusion' debe enviar datos de sÃƒÂ­ntesis."""
        response_mock = Mock()
        response_mock.status_code = 200
        response_mock.json.return_value = {
            "choices": [{"message": {"content": '{"conclusion":"Conclusion generada"}'}}]
        }
        post_mock.return_value = response_mock

        context = {
            "sintesis_conclusion": "No se observa manipulaciÃƒÂ³n",
            "cantidad_observada": "2 personas",
        }
        result = IntegrityService.improve_report_text_with_ai(
            "", "", "texto conclusion", material_context=context, mode='conclusion'
        )

        self.assertEqual(result.get('conclusion'), 'Conclusion generada')
        self.assertEqual(result.get('material_filmico'), '')
        self.assertEqual(result.get('desarrollo'), '')

        _, kwargs = post_mock.call_args
        messages = kwargs.get('json', {}).get('messages', [])
        user_content = json.loads(messages[1].get('content') or '{}')
        self.assertIn('datos_para_conclusion', user_content)
        self.assertEqual(
            user_content['datos_para_conclusion']['sintesis_del_analisis'],
            'No se observa manipulaciÃƒÂ³n'
        )
        self.assertEqual(set(user_content['formato_estricto'].keys()), {'conclusion'})

    def test_mode_rejects_invalid_value(self):
        """El serializer debe rechazar un mode invÃƒÂ¡lido."""
        from .serializers import VideoReportImproveTextSerializer
        serializer = VideoReportImproveTextSerializer(data={
            "material_filmico": "texto",
            "mode": "inventado",
        })
        self.assertFalse(serializer.is_valid())
        self.assertIn('mode', serializer.errors)

