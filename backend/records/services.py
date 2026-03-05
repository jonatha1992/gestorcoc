import hashlib
import os
import base64
import binascii
import json
import re
from threading import Lock
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from pathlib import Path
from django.conf import settings
import requests
import logging
logger = logging.getLogger(__name__)

class IntegrityService:
    _provider_rotation_lock = Lock()
    _provider_rotation_index = 0

    @staticmethod
    def _as_text(value, default=''):
        if value is None:
            return default
        return str(value)

    @staticmethod
    def _as_list(value):
        if isinstance(value, list):
            return value
        return []

    @staticmethod
    def _normalize_hash_algorithms(value):
        allowed = {'sha1', 'sha3', 'sha256', 'sha512'}
        output = []
        for item in IntegrityService._as_list(value):
            normalized = str(item or '').strip().lower()
            if normalized in allowed and normalized not in output:
                output.append(normalized)
        return output

    @staticmethod
    def _hash_algorithm_label(code):
        labels = {
            'sha1': 'SHA-1',
            'sha3': 'SHA-3',
            'sha256': 'SHA-256',
            'sha512': 'SHA-512',
        }
        return labels.get(str(code or '').lower(), str(code or '').upper())

    @staticmethod
    def _vms_authenticity_label(mode):
        labels = {
            'vms_propio': 'autenticación provista por el propio sistema VMS',
            'hash_preventivo': 'autenticación basada en hash preventivo externo',
            'sin_autenticacion': 'sin autenticación técnica declarada',
            'otro': 'autenticación declarada por método alternativo',
        }
        return labels.get(str(mode or '').strip().lower(), 'autenticación no consignada')

    @staticmethod
    def _build_material_filmico_fallback(context):
        ctx = context or {}
        sistema = IntegrityService._as_text(ctx.get('sistema'), 'el sistema de videovigilancia').strip() or 'el sistema de videovigilancia'
        lugar = IntegrityService._as_text(ctx.get('aeropuerto'), '').strip()
        hash_program = IntegrityService._as_text(ctx.get('hash_program'), '').strip()
        hash_algorithms = IntegrityService._normalize_hash_algorithms(ctx.get('hash_algorithms'))
        hash_labels = [IntegrityService._hash_algorithm_label(item) for item in hash_algorithms]
        motivo_sin_hash = IntegrityService._as_text(ctx.get('motivo_sin_hash'), '').strip()
        vms_mode = IntegrityService._as_text(ctx.get('vms_authenticity_mode'), '').strip()
        vms_detail = IntegrityService._as_text(ctx.get('vms_authenticity_detail'), '').strip()
        vms_native_algorithms = IntegrityService._normalize_hash_algorithms(ctx.get('vms_native_hash_algorithms'))
        vms_native_labels = [IntegrityService._hash_algorithm_label(item) for item in vms_native_algorithms]

        lugar_clause = f'desde donde se obtuvo la información en "{lugar}"' if lugar else 'desde donde se obtuvo el material analizado'

        if vms_mode == 'vms_propio':
            if vms_native_labels:
                native_text = ' y '.join(vms_native_labels) if len(vms_native_labels) <= 2 else ', '.join(vms_native_labels[:-1]) + f' y {vms_native_labels[-1]}'
                auth_clause = f'posee como medida de seguridad autenticación provista por el propio sistema "{sistema}", que incorpora algoritmos de hash nativos ({native_text})'
            else:
                auth_clause = f'posee como medida de seguridad autenticación propietaria provista por el propio sistema "{sistema}"'
        elif vms_mode == 'hash_preventivo':
            auth_clause = 'fue sometido a verificación de integridad mediante hash preventivo externo'
        elif vms_mode == 'sin_autenticacion':
            motivo_part = f' ({motivo_sin_hash})' if motivo_sin_hash else ''
            auth_clause = f'no se aplicó método de verificación de integridad al material exportado{motivo_part}'
        elif vms_mode == 'otro':
            detail_part = vms_detail or 'método alternativo no detallado'
            auth_clause = f'se empleó un método alternativo de autenticación: {detail_part}'
        else:
            auth_clause = 'cuenta con medidas de seguridad propias del sistema'

        if hash_labels:
            hash_algorithms_text = ' y '.join(hash_labels) if len(hash_labels) <= 2 else ', '.join(hash_labels[:-1]) + f' y {hash_labels[-1]}'
            hash_program_text = hash_program or 'herramienta de hash'
            hash_clause = (
                f'Asimismo, previamente a su examen, esta instancia procedió a efectuar sobre el material digital '
                f'un hash de seguridad mediante {hash_program_text}, bajo los algoritmos {hash_algorithms_text}, '
                f'con la finalidad de preservar su integridad.'
            )
        elif motivo_sin_hash and vms_mode != 'sin_autenticacion':
            hash_clause = f'Se deja constancia de que no se efectuó hash sobre el material ({motivo_sin_hash}).'
        else:
            hash_clause = ''

        base = (
            f'Es oportuno mencionar que el sistema de videovigilancia denominado "{sistema}", '
            f'{lugar_clause}, {auth_clause}.'
        )
        return f'{base} {hash_clause}'.strip()

    @staticmethod
    def _build_desarrollo_fallback(context):
        ctx = context or {}
        sectores = IntegrityService._as_text(ctx.get('sectores_analizados'), '').strip()
        franja = IntegrityService._as_text(ctx.get('franja_horaria_analizada'), '').strip()
        tiempo = IntegrityService._as_text(ctx.get('tiempo_total_analisis'), '').strip()
        cantidad = IntegrityService._as_text(ctx.get('cantidad_observada'), '').strip()

        partes = []
        if sectores:
            partes.append(f'centrado en los sectores: {sectores}')
        if franja:
            partes.append(f'abarcando la franja horaria {franja}')
        if tiempo:
            partes.append(f'con una duración total de análisis de {tiempo}')

        if partes:
            descripcion = ', '.join(partes)
            base = f'Del análisis visual practicado sobre los registros fílmicos, se deja constancia de que la revisión fue {descripcion}.'
        else:
            base = 'Del análisis visual practicado sobre los registros fílmicos, se deja constancia de la revisión efectuada sobre el material disponible.'

        cantidad_part = f' En cuanto al dato cuantitativo relevante, se consignó: {cantidad}.' if cantidad else ''

        return f'{base}{cantidad_part} La presente descripción se limita al contenido visual observado, sin interpretación pericial.'

    @staticmethod
    def _build_conclusion_fallback(context):
        ctx = context or {}
        sintesis_conclusion = IntegrityService._as_text(ctx.get('sintesis_conclusion'), '').strip()
        cantidad_observada = IntegrityService._as_text(ctx.get('cantidad_observada'), '').strip()

        if sintesis_conclusion:
            return (
                'En virtud del análisis efectuado, y sin apartarse de los extremos objetivamente observables, '
                f'se concluye: {sintesis_conclusion}.'
            )

        if cantidad_observada:
            return (
                'En virtud del análisis efectuado, y dentro de los límites propios de la revisión visual, '
                f'se concluye que la cantidad observada fue de {cantidad_observada}. '
                'No se advierten elementos adicionales a los ya consignados en el desarrollo, '
                'quedando la valoración jurídica sujeta a la autoridad competente.'
            )

        return (
            'En virtud del análisis efectuado sobre el material fílmico, y dentro de los límites propios '
            'de la revisión visual practicada, no se advierten elementos adicionales a los ya consignados '
            'en el desarrollo. La valoración jurídica del material queda sujeta a la autoridad competente.'
        )

    @staticmethod
    def _add_heading_safe(document, text, level=1):
        style_by_level = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
        style_name = style_by_level.get(level)
        if style_name:
            try:
                return document.add_paragraph(text, style=style_name)
            except Exception:
                pass
        paragraph = document.add_paragraph(text)
        if level == 1:
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.size = None
        elif level == 2:
            paragraph.runs[0].bold = True
        return paragraph

    @staticmethod
    def _decode_base64_image(content_base64):
        payload = content_base64.split(',', 1)[1] if ',' in content_base64 else content_base64
        try:
            return base64.b64decode(payload, validate=True)
        except (ValueError, binascii.Error) as exc:
            raise RuntimeError("No se pudo decodificar un fotograma en base64.") from exc

    @staticmethod
    def _prepare_picture_stream(frame):
        raw = IntegrityService._decode_base64_image(frame.get('content_base64', ''))
        mime = frame.get('mime_type')
        if mime != 'image/webp':
            return BytesIO(raw)

        try:
            from PIL import Image
        except ImportError as exc:
            raise RuntimeError("Falta Pillow para procesar imagenes WEBP.") from exc

        try:
            source = BytesIO(raw)
            target = BytesIO()
            with Image.open(source) as image:
                image.convert('RGB').save(target, format='PNG')
            target.seek(0)
            return target
        except Exception as exc:
            raise RuntimeError("No se pudo convertir un fotograma WEBP.") from exc

    @staticmethod
    def _append_frames_annex(document, frames):
        if not frames:
            return

        try:
            from docx.shared import Inches
        except ImportError as exc:
            raise RuntimeError("python-docx no permite insertar imagenes (docx.shared no disponible).") from exc

        ordered = sorted(frames, key=lambda frame: frame.get('order', 0))
        document.add_page_break()
        IntegrityService._add_heading_safe(document, "Anexo de fotogramas", level=1)

        for idx, frame in enumerate(ordered, start=1):
            file_name = frame.get('file_name', f'fotograma_{idx}.jpg')
            description = frame.get('description') or ''
            IntegrityService._add_heading_safe(document, f"Fotograma {idx}: {file_name}", level=2)
            pic_stream = IntegrityService._prepare_picture_stream(frame)
            try:
                document.add_picture(pic_stream, width=Inches(6.2))
            except Exception as exc:
                raise RuntimeError(f"No se pudo insertar el fotograma '{file_name}'.") from exc

            if description:
                document.add_paragraph(f"Descripcion: {description}")

            document.add_paragraph("")

    @staticmethod
    def _normalize_video_payload(payload):
        # Nuevo formato recomendado: {'report_data': {...}, 'frames': [...]}
        if 'report_data' in payload:
            return payload.get('report_data') or {}, payload.get('frames') or []
        # Compatibilidad legacy
        return payload, []

    @staticmethod
    def _replace_text_in_docx(document, replacements):
        def replace_in_paragraph(paragraph):
            # First pass: replace within individual runs (preserves per-run formatting).
            for run in paragraph.runs:
                for src, dst in replacements:
                    if src in run.text:
                        run.text = run.text.replace(src, dst)

            # Second pass: handle placeholders split across multiple runs.
            # Re-read the full paragraph text after the first pass.
            full_text = paragraph.text
            if not any(src in full_text for src, _ in replacements):
                return
            updated = full_text
            for src, dst in replacements:
                updated = updated.replace(src, dst)
            if updated == full_text:
                return
            # Merge all run texts into the first run; clear the rest.
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ''

        for paragraph in document.paragraphs:
            replace_in_paragraph(paragraph)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                replace_in_paragraph(paragraph)
            for paragraph in section.footer.paragraphs:
                replace_in_paragraph(paragraph)

    @staticmethod
    def _extract_json_object(raw_text):
        if isinstance(raw_text, dict):
            return raw_text

        text = str(raw_text or '').strip()
        if not text:
            raise RuntimeError("La API de IA no devolvio contenido para mejorar el texto.")

        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

        match = re.search(r"\{.*\}", text, re.DOTALL)
        if not match:
            raise RuntimeError("No se pudo interpretar la respuesta de la API de IA.")

        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError as exc:
            raise RuntimeError("No se pudo interpretar la respuesta JSON de la API de IA.") from exc

    @staticmethod
    def _vms_mode_to_text(vms_mode, vms_detail=''):
        mapping = {
            'vms_propio': 'VMS propio del sistema de grabacion',
            'hash_preventivo': 'hash preventivo aplicado al material filmico',
            'sin_autenticacion': 'sin medida de autenticacion adicional',
            'otro': vms_detail.strip() if (vms_detail or '').strip() else 'metodo de autenticidad no consignado',
        }
        return mapping.get(str(vms_mode or '').strip(), 'metodo de autenticidad no consignado')

    @staticmethod
    def _build_ai_request_payload(material_filmico, desarrollo, conclusion, model, material_context=None, mode='full'):
        ctx = material_context if isinstance(material_context, dict) else {}

        STYLE_NOTE = (
            "REGLAS DE ESTILO OBLIGATORIAS:\n"
            "1. Usa estilo de acta policial argentina: oraciones completas, voz institucional e impersonal.\n"
            "2. PROHIBIDO inventar hechos, datos, fechas, cantidades o nombres que no esten explicitamente en el contexto provisto.\n"
            "3. Si falta informacion estructurada o no hay datos, ajusta la redaccion para omitirlos de manera fluida en lugar de usar textos como 'no consignado' o 'no determinado'. El texto resultante debe tener total sentido sin estos datos.\n"
            "4. Sin abreviaturas informales. Usar terminos tecnicos precisos.\n"
            "5. Cada apartado debe tener entre 3 y 8 oraciones, claras y directas.\n"
            "6. Devolver UNICAMENTE un JSON valido, sin texto adicional fuera del JSON."
        )

        if mode == 'material_filmico':
            sistema = (ctx.get('sistema') or '').strip()
            aeropuerto = (ctx.get('aeropuerto') or '').strip()
            hash_algorithms = ctx.get('hash_algorithms') or []
            hash_program = (ctx.get('hash_program') or '').strip()
            vms_mode = ctx.get('vms_authenticity_mode') or ''
            vms_detail = ctx.get('vms_authenticity_detail') or ''
            medida = IntegrityService._vms_mode_to_text(vms_mode, vms_detail)
            hash_text = ', '.join(h.upper() for h in hash_algorithms) if hash_algorithms else ''

            system_prompt = (
                "Eres un redactor experto en informes policiales argentinos de analisis de video de CCTV. "
                "Tu tarea es redactar el apartado 'MATERIAL FILMICO ANALIZADO' de un informe oficial. "
                "Este apartado DEBE describir con precision:\n"
                "- El sistema de videovigilancia (VMS/CCTV) utilizado para obtener las grabaciones.\n"
                "- El lugar fisico de donde proviene el material filmico (aeropuerto, dependencia, etc.).\n"
                "- Las medidas de integridad aplicadas: si se uso hash preventivo (indicar algoritmo y programa), "
                "si el VMS posee autenticacion propia, o si no se aplico ningun metodo.\n"
                "- El material debe describirse de forma generica (no inventar camaras ni horarios especificos a menos que se provean).\n\n"
                + STYLE_NOTE
            )
            payload_prompt = {
                "instrucciones": (
                    "Redacta o mejora el texto del apartado de MATERIAL FILMICO ANALIZADO. "
                    "Incorpora los datos del sistema provistos de manera natural. "
                    "Si hay texto_original, mejorarlo. Si no hay, redactar desde cero. "
                    "Devuelve SOLO un JSON valido con la clave 'material_filmico'. Ejemplo: {\"material_filmico\": \"texto...\"}"
                ),
                "formato_estricto": {"material_filmico": "texto redactado"},
                "texto_original": str(material_filmico or ''),
                "datos_del_sistema": {
                    "sistema_cctv": sistema,
                    "lugar_o_aeropuerto": aeropuerto,
                    "algoritmos_hash": hash_text,
                    "programa_hash": hash_program,
                    "medida_autenticidad": medida,
                },
            }

        elif mode == 'desarrollo':
            sectores = (ctx.get('sectores_analizados') or '').strip()
            franja = (ctx.get('franja_horaria_analizada') or '').strip()
            tiempo = (ctx.get('tiempo_total_analisis') or '').strip()
            cantidad = (ctx.get('cantidad_observada') or '').strip()
            caratula = (ctx.get('caratula') or '').strip()
            sistema = (ctx.get('sistema') or '').strip()

            system_prompt = (
                "Eres un redactor experto en informes policiales argentinos de analisis de video de CCTV. "
                "Tu tarea es redactar el apartado 'DESARROLLO' de un informe oficial de analisis de video. "
                "Este apartado DEBE narrar de forma objetiva y cronologica lo observado.\n"
                "- Limitar la descripcion a lo objetivamente observable en el material filmico.\n\n"
                + STYLE_NOTE
            )
            payload_prompt = {
                "instrucciones": (
                    "Redacta o mejora el texto del apartado DESARROLLO usando los datos provistos. "
                    "Si hay texto_original, mejorar su redaccion manteniendo la informacion. "
                    "Si texto_original esta vacio, redactar desde cero. "
                    "Devuelve SOLO un JSON valido con la clave 'desarrollo'. Ejemplo: {\"desarrollo\": \"texto...\"}"
                ),
                "formato_estricto": {"desarrollo": "texto redactado"},
                "texto_original": str(desarrollo or ''),
                "datos_del_analisis": {
                    "sistema_cctv": sistema,
                    "sectores_analizados": sectores,
                    "franja_horaria": franja,
                    "tiempo_total_analisis": tiempo,
                    "cantidad_observada": cantidad,
                    "caratula_causa": caratula,
                },
            }

        elif mode == 'conclusion':
            sintesis = (ctx.get('sintesis_conclusion') or '').strip()
            cantidad = (ctx.get('cantidad_observada') or '').strip()
            caratula = (ctx.get('caratula') or '').strip()

            system_prompt = (
                "Eres un redactor experto en informes policiales argentinos de analisis de video de CCTV. "
                "Tu tarea es redactar el apartado 'CONCLUSION' de un informe oficial de analisis de video. "
                "La conclusion DEBE sintetizar el resultado del analisis visual e institucional.\n"
                "- Si se provee una sintesis_del_analisis, priorizarla como eje central.\n"
                "- Ser un parrafo de cierre formal.\n\n"
                + STYLE_NOTE
            )
            payload_prompt = {
                "instrucciones": (
                    "Redacta o mejora la CONCLUSION. "
                    "Devuelve SOLO un JSON valido con la clave 'conclusion'. Ejemplo: {\"conclusion\": \"texto...\"}"
                ),
                "formato_estricto": {"conclusion": "texto redactado"},
                "texto_original": str(conclusion or ''),
                "datos_para_conclusion": {
                    "sintesis_del_analisis": sintesis,
                    "cantidad_observada": cantidad,
                    "caratula_causa": caratula,
                },
            }

        else:  # mode == 'full'
            system_prompt = (
                "Eres un redactor experto en informes policiales argentinos de analisis de video de CCTV. "
                "Tu tarea es mejorar o redactar desde cero los tres apartados narrativos de un informe oficial: "
                "'material filmico analizado', 'desarrollo' y 'conclusion'.\n"
                "Cada apartado debe ser coherente con el anterior y mantener un tono institucional uniforme. " + STYLE_NOTE
            )
            payload_prompt = {
                "instrucciones": (
                    "Mejora y/o redacta desde cero los tres apartados usando el contexto estructurado provisto. "
                    "Se sumamente inteligente: los datos que faltan deben excluirse fluidamente. Crea un texto generico pero realista en estilo policial si la mayor parte de la informacion falta y no hay textos originales."
                    "PROHIBIDO inventar hechos concretos no provistos. "
                    "Devuelve SOLO un JSON valido con TRES CLAVES obligatorias. "
                    "Ejemplo: {\"material_filmico\": \"...\", \"desarrollo\": \"...\", \"conclusion\": \"...\"}"
                ),
                "formato_estricto": {
                    "material_filmico": "texto mejorado o redactado desde cero",
                    "desarrollo": "texto mejorado o redactado desde cero",
                    "conclusion": "texto mejorado o redactado desde cero",
                },
                "material_filmico_original": str(material_filmico or ''),
                "desarrollo_original": str(desarrollo or ''),
                "conclusion_original": str(conclusion or ''),
                "material_context": ctx,
            }

        return {
            "model": model,
            "temperature": 0.2,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": json.dumps(payload_prompt, ensure_ascii=False)},
            ],
            "response_format": {"type": "json_object"},
        }

    @staticmethod
    def _extract_ai_error_detail(response):
        detail = ""
        try:
            error_payload = response.json()
            error_node = error_payload.get('error', {})
            if isinstance(error_node, dict):
                detail = str(error_node.get('message', '') or '')
            elif error_node:
                detail = str(error_node)
        except Exception:
            detail = response.text[:200]
        return detail

    @staticmethod
    def _is_quota_or_tokens_error(provider, status_code, detail):
        message = str(detail or '').lower()

        if provider == 'openrouter':
            return (
                status_code in (402, 429)
                or 'insufficient credits' in message
                or 'out of credits' in message
                or 'quota' in message
                or 'token' in message
            )

        if provider == 'gemini':
            return (
                status_code == 429
                or 'resource_exhausted' in message
                or 'resource has been exhausted' in message
                or 'quota' in message
                or 'token' in message
            )

        if provider == 'groq':
            return (
                status_code == 429
                or 'rate limit' in message
                or 'quota' in message
                or 'token' in message
            )
            
        if provider == 'ollama':
            # Locales no suelen tener cuota, pero por si acaso.
            return False

        return False

    @staticmethod
    def _get_ai_provider_chain(custom_api_key='', preferred_provider=''):
        provider_order_raw = str(
            getattr(settings, 'AI_TEXT_PROVIDER_ORDER', 'ollama,gemini,openrouter,groq') or ''
        ).strip()
        provider_order = [item.strip().lower() for item in provider_order_raw.split(',') if item.strip()]
        if not provider_order:
            provider_order = ['ollama', 'gemini', 'openrouter', 'groq']

        if preferred_provider:
            preferred = str(preferred_provider).strip().lower()
            if preferred in provider_order:
                provider_order.remove(preferred)
                provider_order.insert(0, preferred)
            # If not in the pre-configured order, just add it at the top
            elif preferred in ['ollama', 'gemini', 'openrouter', 'groq']:
                provider_order.insert(0, preferred)

        selection_mode = str(
            getattr(settings, 'AI_TEXT_PROVIDER_SELECTION', 'ordered') or 'ordered'
        ).strip().lower()
        if selection_mode == 'round_robin' and len(provider_order) > 1 and not preferred_provider:
            with IntegrityService._provider_rotation_lock:
                start_idx = IntegrityService._provider_rotation_index % len(provider_order)
                IntegrityService._provider_rotation_index += 1
            provider_order = provider_order[start_idx:] + provider_order[:start_idx]

        providers = {
            'gemini': {
                'name': 'gemini',
                'api_key': str(getattr(settings, 'GEMINI_API_KEY', '') or '').strip(),
                'api_url': str(getattr(settings, 'AI_TEXT_GEMINI_API_URL', '') or '').strip(),
                'model': str(getattr(settings, 'AI_TEXT_GEMINI_MODEL', '') or '').strip(),
            },
            'openrouter': {
                'name': 'openrouter',
                'api_key': str(
                    getattr(settings, 'OPEN_ROUTER_API_KEY', '') or getattr(settings, 'AI_TEXT_API_KEY', '') or ''
                ).strip(),
                'api_url': str(
                    getattr(settings, 'AI_TEXT_OPENROUTER_API_URL', '') or getattr(settings, 'AI_TEXT_API_URL', '') or ''
                ).strip(),
                'model': str(
                    getattr(settings, 'AI_TEXT_OPENROUTER_MODEL', '') or getattr(settings, 'AI_TEXT_MODEL', '') or ''
                ).strip(),
            },
            'groq': {
                'name': 'groq',
                'api_key': str(getattr(settings, 'GROQ_API_KEY', '') or '').strip(),
                'api_url': str(getattr(settings, 'AI_TEXT_GROQ_API_URL', '') or '').strip(),
                'model': str(getattr(settings, 'AI_TEXT_GROQ_MODEL', '') or '').strip(),
            },
            'ollama': {
                'name': 'ollama',
                'api_key': str(getattr(settings, 'OLLAMA_API_KEY', 'ollama') or 'ollama').strip(),  # Ollama doesn't strictly need it, but for standardized headers
                'api_url': str(getattr(settings, 'AI_TEXT_OLLAMA_API_URL', 'http://localhost:11434/v1/chat/completions') or '').strip(),
                'model': str(getattr(settings, 'AI_TEXT_OLLAMA_MODEL', 'llama3.2') or '').strip(),
            },
        }

        chain = []
        manual_key = str(custom_api_key or '').strip()

        for index, provider_name in enumerate(provider_order):
            if provider_name not in providers:
                continue

            provider = dict(providers[provider_name])
            if index == 0 and manual_key:
                provider['api_key'] = manual_key

            if not provider.get('api_key') or not provider.get('api_url') or not provider.get('model'):
                continue

            chain.append(provider)

        return chain

    @staticmethod
    def improve_report_text_with_ai(material_filmico, desarrollo, conclusion, custom_api_key="", material_context=None, mode='full', preferred_provider=""):
        timeout_seconds = int(getattr(settings, 'AI_TEXT_TIMEOUT_SECONDS', 45))
        fallback_mode = str(getattr(settings, 'AI_TEXT_FALLBACK_MODE', 'quota_only') or 'quota_only').strip().lower()
        providers = IntegrityService._get_ai_provider_chain(custom_api_key, preferred_provider)

        if not providers:
            raise RuntimeError(
                "No se configuro AI_TEXT_PROVIDER_KEYS "
                "(GEMINI_API_KEY, OPEN_ROUTER_API_KEY o GROQ_API_KEY) "
                "y no se ingreso ninguna manualmente."
            )

        for index, provider in enumerate(providers):
            request_payload = IntegrityService._build_ai_request_payload(
                material_filmico,
                desarrollo,
                conclusion,
                provider['model'],
                material_context=material_context,
                mode=mode,
            )
            headers = {
                "Authorization": f"Bearer {provider['api_key']}",
                "Content-Type": "application/json",
            }

            try:
                response = requests.post(
                    url=provider['api_url'],
                    headers=headers,
                    json=request_payload,
                    timeout=timeout_seconds,
                )
            except requests.RequestException as exc:
                raise RuntimeError("No se pudo conectar con la API de IA.") from exc

            if response.status_code >= 400:
                detail = IntegrityService._extract_ai_error_detail(response)
                if (
                    fallback_mode == 'quota_only'
                    and IntegrityService._is_quota_or_tokens_error(
                        provider['name'],
                        response.status_code,
                        detail,
                    )
                ):
                    if index < len(providers) - 1:
                        continue
                    raise RuntimeError("Todos los proveedores de IA alcanzaron su limite de cuota/tokens.")

                suffix = f" Detalle: {detail}" if detail else ""
                raise RuntimeError(
                    f"La API de IA devolvio un error ({response.status_code}).{suffix}"
                )

            try:
                response_payload = response.json()
                choices = response_payload.get('choices') or []
                if not choices:
                    raise RuntimeError("La API de IA no devolvio opciones de respuesta.")

                message = choices[0].get('message') if isinstance(choices[0], dict) else None
                content = message.get('content') if isinstance(message, dict) else None
                if isinstance(content, list):
                    parts = []
                    for item in content:
                        if isinstance(item, dict):
                            parts.append(str(item.get('text', '') or ''))
                        else:
                            parts.append(str(item))
                    content = ''.join(parts)

                parsed_content = IntegrityService._extract_json_object(content)

                orig_mf = str(material_filmico or '').strip()
                orig_des = str(desarrollo or '').strip()
                orig_con = str(conclusion or '').strip()

                if mode == 'material_filmico':
                    improved_material_filmico = IntegrityService._as_text(
                        parsed_content.get('material_filmico'), orig_mf
                    ).strip()
                    improved_desarrollo = orig_des
                    improved_conclusion = orig_con
                elif mode == 'desarrollo':
                    improved_material_filmico = orig_mf
                    improved_desarrollo = IntegrityService._as_text(
                        parsed_content.get('desarrollo'), orig_des
                    ).strip()
                    improved_conclusion = orig_con
                elif mode == 'conclusion':
                    improved_material_filmico = orig_mf
                    improved_desarrollo = orig_des
                    improved_conclusion = IntegrityService._as_text(
                        parsed_content.get('conclusion'), orig_con
                    ).strip()
                else:
                    improved_material_filmico = IntegrityService._as_text(
                        parsed_content.get('material_filmico'), orig_mf
                    ).strip()
                    improved_desarrollo = IntegrityService._as_text(
                        parsed_content.get('desarrollo'), orig_des
                    ).strip()
                    improved_conclusion = IntegrityService._as_text(
                        parsed_content.get('conclusion'), orig_con
                    ).strip()

                ai_applied = any([
                    improved_material_filmico != orig_mf,
                    improved_desarrollo != orig_des,
                    improved_conclusion != orig_con,
                ])
                usage = response_payload.get('usage') or {}
                tokens_in    = usage.get('prompt_tokens', 0)
                tokens_out   = usage.get('completion_tokens', 0)
                tokens_total = usage.get('total_tokens', tokens_in + tokens_out)
                logger.info(
                    "[AI-TOKENS] provider=%s model=%s in=%d out=%d total=%d",
                    provider['name'], provider['model'], tokens_in, tokens_out, tokens_total
                )
                try:
                    from records.models import AIUsageLog
                    AIUsageLog.objects.create(
                        provider=provider['name'],
                        model_name=provider['model'],
                        endpoint='improve_text',
                        tokens_in=tokens_in,
                        tokens_out=tokens_out,
                        tokens_total=tokens_total,
                        success=True,
                    )
                except Exception:
                    logger.exception("No se pudo guardar AIUsageLog")
                return {
                    'material_filmico': improved_material_filmico,
                    'desarrollo': improved_desarrollo,
                    'conclusion': improved_conclusion,
                    'ai_applied': ai_applied,
                }
            except RuntimeError:
                raise
            except Exception as exc:
                raise RuntimeError("Respuesta invalida al mejorar texto con IA.") from exc

        raise RuntimeError("No se encontro un proveedor de IA utilizable.")

    @staticmethod
    def generate_video_analysis_docx(payload):
        """
        Genera un informe DOCX usando data/Modelo Informe.docx como base.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx no esta instalado") from exc

        template_path = Path(__file__).resolve().parents[3] / "data" / "Modelo Informe.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"No se encontro plantilla: {template_path}")

        report_data, frames = IntegrityService._normalize_video_payload(payload or {})

        report_date = report_data.get("report_date") or datetime.now().strftime("%Y-%m-%d")
        try:
            if isinstance(report_date, datetime):
                parsed = report_date
            else:
                parsed = datetime.strptime(str(report_date), "%Y-%m-%d")
        except ValueError:
            parsed = datetime.now()

        month_names = {
            1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
            5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
            9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
        }
        month_name = month_names.get(parsed.month, "enero")
        year = str(parsed.year)

        destinatarios = IntegrityService._as_text(report_data.get("destinatarios"), "URSA I - Jefe")
        tipo_informe = IntegrityService._as_text(report_data.get("tipo_informe"), "IAV - Informe Analisis de Video")
        numero_informe = IntegrityService._as_text(report_data.get("numero_informe"), f"001/{parsed.year}")
        grado = IntegrityService._as_text(report_data.get("grado"), "Oficial Mayor")
        operador = IntegrityService._as_text(report_data.get("operador"), "Operador")
        lup = IntegrityService._as_text(report_data.get("lup"), "0000")
        sistema = IntegrityService._as_text(report_data.get("sistema"), "MILESTONE")
        cantidad_observada = IntegrityService._as_text(report_data.get("cantidad_observada"), "").strip()
        sectores_analizados = IntegrityService._as_text(report_data.get("sectores_analizados"), "").strip()
        franja_horaria_analizada = IntegrityService._as_text(report_data.get("franja_horaria_analizada"), "").strip()
        tiempo_total_analisis = IntegrityService._as_text(report_data.get("tiempo_total_analisis"), "").strip()
        sintesis_conclusion = IntegrityService._as_text(report_data.get("sintesis_conclusion"), "").strip()
        hash_algorithms = IntegrityService._normalize_hash_algorithms(report_data.get("hash_algorithms"))
        hash_program = IntegrityService._as_text(report_data.get("hash_program"), "").strip()
        medida_seguridad_interna = IntegrityService._as_text(report_data.get("medida_seguridad_interna"), "").strip()
        vms_authenticity_mode = IntegrityService._as_text(report_data.get("vms_authenticity_mode"), "").strip()
        vms_authenticity_detail = IntegrityService._as_text(report_data.get("vms_authenticity_detail"), "").strip()
        prevencion_sumaria = IntegrityService._as_text(report_data.get("prevencion_sumaria"), "000BAR/2026")
        caratula = IntegrityService._as_text(report_data.get("caratula"), "DENUNCIA S/ PRESUNTO HURTO")
        unidad = IntegrityService._as_text(report_data.get("unidad"), "Coordinación Regional de Video Seguridad I del Este")
        material_filmico = IntegrityService._as_text(report_data.get("material_filmico"), "material fílmico")
        fiscalia = IntegrityService._as_text(report_data.get("fiscalia"), "").strip()
        fiscal = IntegrityService._as_text(report_data.get("fiscal"), "").strip()
        fiscalia_doc = fiscalia or "Fiscalia / Juzgado no consignado"
        fiscal_doc = fiscal or "Autoridad interviniente no consignada"
        denunciante = IntegrityService._as_text(report_data.get("denunciante"), "Denunciante")
        vuelo = IntegrityService._as_text(report_data.get("vuelo"), "---")
        empresa_aerea = IntegrityService._as_text(report_data.get("empresa_aerea"), "---")
        destino = IntegrityService._as_text(report_data.get("destino"), "---")
        fecha_hecho = IntegrityService._as_text(report_data.get("fecha_hecho"), "---")
        aeropuerto = IntegrityService._as_text(report_data.get("aeropuerto"), "---")
        objeto_denunciado = IntegrityService._as_text(report_data.get("objeto_denunciado"), "objeto denunciado")
        conclusion = IntegrityService._as_text(report_data.get("conclusion"), "")
        desarrollo = IntegrityService._as_text(report_data.get("desarrollo"), "")
        firma = IntegrityService._as_text(report_data.get("firma"), "Coordinador CReV I DEL ESTE")

        if not material_filmico.strip():
            material_filmico = IntegrityService._build_material_filmico_fallback({
                "sistema": sistema,
                "aeropuerto": aeropuerto,
                "cantidad_observada": cantidad_observada,
                "sectores_analizados": sectores_analizados,
                "franja_horaria_analizada": franja_horaria_analizada,
                "tiempo_total_analisis": tiempo_total_analisis,
                "hash_algorithms": hash_algorithms,
                "hash_program": hash_program,
                "medida_seguridad_interna": medida_seguridad_interna,
                "vms_authenticity_mode": vms_authenticity_mode,
                "vms_authenticity_detail": vms_authenticity_detail,
            })

        if not desarrollo.strip():
            desarrollo = IntegrityService._build_desarrollo_fallback({
                "sectores_analizados": sectores_analizados,
                "franja_horaria_analizada": franja_horaria_analizada,
                "tiempo_total_analisis": tiempo_total_analisis,
                "cantidad_observada": cantidad_observada,
            })

        if not conclusion.strip():
            conclusion = IntegrityService._build_conclusion_fallback({
                "sintesis_conclusion": sintesis_conclusion,
                "cantidad_observada": cantidad_observada,
            })

        op_info = f"{grado} {operador}, LUP: {lup}"

        # Curly-quote characters used by Word as default punctuation
        LQ = "\u201c"  # “ left double quotation mark
        RQ = "\u201d"  # ” right double quotation mark

        # Build hash algorithm text for paragraph [9] replacement
        if hash_algorithms:
            algo_labels = [IntegrityService._hash_algorithm_label(a) for a in hash_algorithms]
            if len(algo_labels) == 1:
                hash_algo_text = f"{algo_labels[0]} respectivamente"
            elif len(algo_labels) == 2:
                hash_algo_text = f"{algo_labels[0]} y {algo_labels[1]} respectivamente"
            else:
                hash_algo_text = f"{chr(44).join(algo_labels[:-1])} y {algo_labels[-1]} respectivamente"
        else:
            hash_algo_text = "algoritmos no consignados respectivamente"

        medida_doc = medida_seguridad_interna or "no consignada"

        replacements = [
            # Header / title fields (may appear in headers or title blocks)
            ("Fecha: Enero de 2026.", f"Fecha: {month_name.capitalize()} de {year}."),
            ("URSA I - Jefe", destinatarios),
            ("IAV - Informe Análisis de Video    /2026", f"{tipo_informe} {numero_informe}"),
            ("IAV - Informe Analisis de Video    /2026", f"{tipo_informe} {numero_informe}"),
            # Operative info -- paragraph [6]
            ("Oficial Mayor XXX, LUP: XXX", op_info),
            # sistema: template uses Word curly-quote delimiters around MILESTONE
            (f"denominado {LQ}MILESTONE{RQ}", f"denominado {LQ}{sistema}{RQ}"),
            ('denominado \"MILESTONE\"', f'denominado \"{sistema}\"'),
            ("Prevención Sumaria 003BAR/2026", f"Prevencion Sumaria {prevencion_sumaria}"),
            # Caratula: template has "xxxxxxx" placeholder (not the form default value)
            (f"caratulada {LQ}xxxxxxx{RQ}", f"caratulada {LQ}{caratula}{RQ}"),
            ('caratulada \"xxxxxxx\"', f'caratulada \"{caratula}\"'),
            ("Fiscalía Nro. 02 de San Carlos de Bariloche, Provincia de Rio Negro", fiscalia_doc),
            ("Dr. INTI ISLA", fiscal_doc),
            ("Sr. PONZO Osvaldo", f"Sr. {denunciante}"),
            ("vuelo WJ 3045", f"vuelo {vuelo}"),
            # empresa_aerea: match the specific airline name used in the template
            ("empresa aerocomercial Jet Smart", empresa_aerea if empresa_aerea and empresa_aerea != "---" else "empresa aerocomercial Jet Smart"),
            ("RIVER PLATE", objeto_denunciado),
            # Use full context strings to avoid replacing common words mid-sentence
            ("con destino a la ciudad de San Carlos de Bariloche", f"con destino a {destino}" if destino and destino != "---" else "con destino a la ciudad de San Carlos de Bariloche"),
            ("el día cuatro de enero del año en curso", f"el día {fecha_hecho}" if fecha_hecho and fecha_hecho != "---" else "el día cuatro de enero del año en curso"),
            # Jurisdiction / unit -- paragraph [6]
            ("Coordinación Regional de Video Seguridad I del Este", unidad),
            ("asiento de la Coordinación Regional de Video Seguridad I del Este", f"de la {unidad}"),
            ("Aeropuerto Internacional Mtro. Pistarini", aeropuerto),
            ("Aeroparque Jorge Newbery", aeropuerto),
            # Paragraph [9]: security measure, hash program and algorithms
            ("medida de seguridad Interna xxxxxxx,", f"medida de seguridad Interna {medida_doc},"),
            ("HASH MI FYLE", hash_program if hash_program else "programa no consignado"),
            ("SHA - 256 y SHA - 512 respectivamente", hash_algo_text),
            # Signature
            ("Coordinador CReV I DEL ESTE", firma),
        ]

        document = Document(str(template_path))
        IntegrityService._replace_text_in_docx(document, replacements)

        # Populate content tables (T0=Material Filmico, T1=Desarrollo, T2=Conclusion)
        tables = document.tables
        if len(tables) > 0:
            row = tables[0].add_row()
            row.cells[0].paragraphs[0].add_run(material_filmico)
        if len(tables) > 1:
            row = tables[1].add_row()
            row.cells[0].paragraphs[0].add_run(desarrollo)
        if len(tables) > 2:
            row = tables[2].add_row()
            row.cells[0].paragraphs[0].add_run(conclusion)

        IntegrityService._append_frames_annex(document, frames)

        out = BytesIO()
        document.save(out)
        out.seek(0)
        return out, f"informe_analisis_video_{parsed.strftime('%Y%m%d')}.docx"

    @staticmethod
    def _get_hasher(algorithm='sha256'):
        if algorithm == 'sha256':
            return hashlib.sha256()
        if algorithm == 'sha512':
            return hashlib.sha512()
        if algorithm == 'sha3':
            return hashlib.sha3_256()
        raise ValueError(f"Algoritmo no soportado: {algorithm}")

    @staticmethod
    def calculate_file_hash(file_obj, algorithm='sha256'):
        """
        Calcula el hash de un archivo usando el algoritmo especificado.
        
        Args:
            file_obj: Objeto de archivo (Django UploadedFile o file-like object)
            algorithm: 'sha256', 'sha512' o 'sha3' (default: sha256)
        
        Returns:
            str: Hash hexadecimal del archivo
        """
        hasher = IntegrityService._get_hasher(algorithm)

        # Leer archivo en chunks para evitar problemas de memoria
        for chunk in file_obj.chunks():
            hasher.update(chunk)
            
        return hasher.hexdigest()
    
    @staticmethod
    def verify_existing_backup(backup_path, expected_hash, algorithm='sha256'):
        """
        Verifica la integridad de un archivo de backup existente.
        
        Args:
            backup_path: Ruta completa al archivo de backup
            expected_hash: Hash esperado para comparación
            algorithm: Algoritmo de hash utilizado
        
        Returns:
            dict: {'is_valid': bool, 'calculated_hash': str, 'message': str}
        """
        if not os.path.exists(backup_path):
            return {
                'is_valid': False,
                'calculated_hash': None,
                'message': f'El archivo no existe en la ruta: {backup_path}'
            }
        
        try:
            hasher = IntegrityService._get_hasher(algorithm)
            
            # Leer archivo en chunks
            with open(backup_path, 'rb') as f:
                while chunk := f.read(8192):
                    hasher.update(chunk)
            
            calculated_hash = hasher.hexdigest()
            is_valid = calculated_hash == expected_hash
            
            return {
                'is_valid': is_valid,
                'calculated_hash': calculated_hash,
                'message': 'Integridad verificada correctamente' if is_valid else 'El hash no coincide - archivo modificado'
            }
        except Exception as e:
            return {
                'is_valid': False,
                'calculated_hash': None,
                'message': f'Error al verificar archivo: {str(e)}'
            }

    @staticmethod
    def generate_integrity_pdf(file_name, file_size, hashes):
        """
        Genera un reporte PDF básico con detalles del archivo y hashes calculados.
        
        Args:
            file_name: Nombre del archivo
            file_size: Tamaño en bytes
            hashes: dict {'SHA256': '...', 'SHA512': '...', 'SHA-3': '...'}
        
        Returns:
            BytesIO: Buffer con el PDF generado
        """
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # Header
        p.setFont("Helvetica-Bold", 18)
        p.drawString(50, height - 50, "Reporte de Integridad de Archivo")
        
        p.setFont("Helvetica", 10)
        p.drawString(50, height - 70, f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        p.line(50, height - 80, width - 50, height - 80)

        # File Details
        y = height - 120
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, "Detalles del Archivo")
        y -= 20
        p.setFont("Helvetica", 12)
        p.drawString(70, y, f"Nombre: {file_name}")
        y -= 20
        p.drawString(70, y, f"Tamaño: {file_size} bytes")
        
        y -= 40
        p.line(50, y, width - 50, y)
        y -= 30

        # Hashes
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, "Verificación de Integridad (Hashes)")
        y -= 30

        for algo, hash_val in hashes.items():
            p.setFont("Helvetica-Bold", 11)
            p.drawString(70, y, f"{algo}:")
            p.setFont("Courier", 10)
            p.drawString(130, y, hash_val)
            y -= 25

        # Footer
        p.setFont("Helvetica-Oblique", 8)
        p.drawString(50, 50, "Este documento es un reporte generado automáticamente por el sistema GestorCOC.")
        p.drawString(50, 40, "No garantiza la autenticidad del archivo si este es modificado posteriormente.")

        p.showPage()
        p.save()
        
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_integrity_summary_pdf(entries):
        """
        Genera un PDF resumen en orientación landscape con todos los archivos
        hasheados en la sesión UI. El hash se muestra completo.

        Args:
            entries: list[dict] con name, size, algorithm, hash, time_ms

        Returns:
            BytesIO: Buffer con el PDF generado
        """
        buffer = BytesIO()
        page_size = landscape(A4)
        p = canvas.Canvas(buffer, pagesize=page_size)
        width, height = page_size

        # Márgenes y ancho útil
        margin = 40
        usable_width = width - (margin * 2)

        def draw_header():
            p.setFont("Helvetica-Bold", 16)
            p.drawString(margin, height - 35, "Resumen de Verificación de Integridad")
            p.setFont("Helvetica", 9)
            p.drawString(margin, height - 52, f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}   |   Total archivos: {len(entries)}")
            p.line(margin, height - 60, width - margin, height - 60)

        def draw_column_headers(y_pos):
            p.setFillColorRGB(0.11, 0.22, 0.37)  # Azul oscuro
            p.rect(margin, y_pos - 4, usable_width, 18, fill=1, stroke=0)
            p.setFillColorRGB(1, 1, 1)
            p.setFont("Helvetica-Bold", 8)
            p.drawString(margin + 4,      y_pos + 2, "#")
            p.drawString(margin + 20,     y_pos + 2, "Nombre")
            p.drawString(margin + 170,    y_pos + 2, "Tipo")
            p.drawString(margin + 210,    y_pos + 2, "Tamaño")
            p.drawString(margin + 260,    y_pos + 2, "Algoritmo")
            p.drawString(margin + 310,    y_pos + 2, "Tiempo")
            p.drawString(margin + 360,    y_pos + 2, "Hash completo")
            p.setFillColorRGB(0, 0, 0)
            return y_pos - 20

        draw_header()
        y = height - 80
        y = draw_column_headers(y)

        # Ancho máximo del campo hash en puntos (desde x=400 hasta el margen derecho)
        hash_x = margin + 360
        hash_max_width = width - margin - hash_x
        # SHA-256 = 64 chars, SHA-512 = 128 chars, SHA-3 = 64 chars
        # Con Courier 7, un car es ~4.2pt → caben ~hash_max_width/4.2 chars por línea
        chars_per_line = int(hash_max_width / 4.2)

        row_bg_even = (0.97, 0.98, 1.0)
        row_bg_odd  = (1.0, 1.0, 1.0)

        for idx, entry in enumerate(entries, start=1):
            name       = str(entry.get("name", "N/A"))
            ext_type   = str(entry.get("type", "N/A"))
            size_str   = str(entry.get("size_str", "N/A"))
            algorithm  = str(entry.get("algorithm", "N/A"))
            time_ms    = str(entry.get("time_ms", "N/A")).strip()
            hash_value = str(entry.get("hash", "N/A"))

            # Quitar posibles dobles ms si la UI los envía por error
            if time_ms.endswith(" ms ms"):
                time_ms = time_ms[:-3]

            # Partir el hash en líneas si es muy largo
            hash_lines = [hash_value[i:i + chars_per_line] for i in range(0, len(hash_value), chars_per_line)] if hash_value != "N/A" else [hash_value]
            row_height = max(14, len(hash_lines) * 10 + 4)

            if y - row_height < margin + 20:
                p.showPage()
                draw_header()
                y = height - 80
                y = draw_column_headers(y)

            # Fondo alternado
            bg = row_bg_even if idx % 2 == 0 else row_bg_odd
            p.setFillColorRGB(*bg)
            p.rect(margin, y - row_height + 12, usable_width, row_height, fill=1, stroke=0)
            p.setFillColorRGB(0, 0, 0)

            # Datos de la fila
            row_top = y
            p.setFont("Helvetica", 8)
            p.drawString(margin + 4,   row_top, str(idx))
            p.drawString(margin + 20,  row_top, name[:25] if len(name) > 25 else name)
            p.drawString(margin + 170, row_top, ext_type[:7] if len(ext_type) > 7 else ext_type)
            p.drawString(margin + 210, row_top, size_str)
            p.drawString(margin + 260, row_top, algorithm)
            p.drawString(margin + 310, row_top, str(time_ms))

            # Hash completo en líneas
            p.setFont("Courier", 7)
            p.setFillColorRGB(0.21, 0.19, 0.64)  # Índigo
            for line_idx, hash_line in enumerate(hash_lines):
                p.drawString(hash_x, row_top - (line_idx * 9), hash_line)
            p.setFillColorRGB(0, 0, 0)

            # Línea separadora
            p.setStrokeColorRGB(0.88, 0.91, 0.94)
            p.line(margin, y - row_height + 12, width - margin, y - row_height + 12)
            p.setStrokeColorRGB(0, 0, 0)

            y -= row_height

        # Footer
        p.setFont("Helvetica-Oblique", 7)
        p.setFillColorRGB(0.58, 0.64, 0.72)
        p.drawString(margin, margin - 10, "Documento generado automáticamente por GestorCOC. Los hashes son calculados por el cliente y no garantizan la cadena de custodia si los archivos son modificados posteriormente.")
        p.setFillColorRGB(0, 0, 0)

        p.showPage()
        p.save()
        buffer.seek(0)
        return buffer

    
    @staticmethod
    def generate_verification_report(film_record):
        """
        Genera un reporte completo de verificación para un FilmRecord.
        Incluye todos los datos del registro, verificación CREV y hash de integridad.
        
        Args:
            film_record: Instancia de FilmRecord
        
        Returns:
            BytesIO: Buffer con el PDF de certificación
        """
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # Header
        p.setFont("Helvetica-Bold", 20)
        p.drawString(50, height - 50, "CERTIFICADO DE INTEGRIDAD")
        p.drawString(50, height - 75, "Registro Fílmico - GestorCOC")
        
        p.setFont("Helvetica", 9)
        p.drawString(50, height - 95, f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        p.line(50, height - 105, width - 50, height - 105)

        # Información del Registro
        y = height - 140
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, y, f"Registro N° {film_record.id}")
        y -= 30
        
        # Datos principales
        p.setFont("Helvetica-Bold", 11)
        p.drawString(50, y, "Información de la Solicitud:")
        y -= 20
        p.setFont("Helvetica", 10)
        
        fields = [
            ("Nº Asunto", film_record.issue_number),
            ("Nº Causa Judicial", film_record.judicial_case_number),
            ("Solicitante", film_record.requester),
            ("Fecha de Ingreso", film_record.entry_date.strftime('%d/%m/%Y') if film_record.entry_date else 'N/A'),
            ("Cámara", film_record.camera.name if film_record.camera else 'N/A'),
            ("Período", f"{film_record.start_time.strftime('%d/%m/%Y %H:%M')} - {film_record.end_time.strftime('%d/%m/%Y %H:%M')}"),
        ]
        
        for label, value in fields:
            if value:
                p.drawString(70, y, f"{label}: {value}")
                y -= 18
        
        # Información de Backup
        y -= 20
        p.setFont("Helvetica-Bold", 11)
        p.drawString(50, y, "Información de Backup:")
        y -= 20
        p.setFont("Helvetica", 10)
        p.drawString(70, y, f"Ruta: {film_record.backup_path or 'No especificada'}")
        y -= 18
        p.drawString(70, y, f"Tamaño: {film_record.file_size or 'N/A'} bytes")
        
        # Hash de Integridad
        y -= 30
        p.setFont("Helvetica-Bold", 11)
        p.drawString(50, y, "Hash de Integridad (SHA-256):")
        y -= 20
        p.setFont("Courier", 8)
        p.drawString(70, y, film_record.file_hash or 'No calculado')
        
        # Verificación CREV
        y -= 40
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, "VERIFICACIÓN CREV")
        p.line(50, y - 5, width - 50, y - 5)
        y -= 25
        p.setFont("Helvetica", 10)
        
        if film_record.verified_by_crev:
            p.drawString(70, y, f"Verificado por: {film_record.verified_by_crev.last_name}, {film_record.verified_by_crev.first_name}")
            y -= 18
            p.drawString(70, y, f"Fecha de Verificación: {film_record.verification_date.strftime('%d/%m/%Y %H:%M:%S')}")
            y -= 18
            p.setFont("Helvetica-Bold", 10)
            p.setFillColorRGB(0, 0.5, 0)  # Verde
            p.drawString(70, y, "✓ REGISTRO CERTIFICADO Y BLOQUEADO PARA EDICIÓN")
        else:
            p.setFillColorRGB(0.8, 0, 0)  # Rojo
            p.drawString(70, y, "✗ PENDIENTE DE VERIFICACIÓN CREV")
        
        # Footer
        p.setFillColorRGB(0, 0, 0)  # Negro
        p.setFont("Helvetica-Oblique", 8)
        p.drawString(50, 50, "Este certificado es válido solo si el hash coincide con el archivo de backup.")
        p.drawString(50, 40, "Sistema GestorCOC - Centro de Operaciones y Control")

        p.showPage()
        p.save()
        
        buffer.seek(0)
        return buffer
