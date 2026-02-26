import hashlib
import os
import base64
import binascii
import json
import re
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from pathlib import Path
from django.conf import settings
import requests

class IntegrityService:
    @staticmethod
    def _as_text(value, default=''):
        if value is None:
            return default
        return str(value)

    @staticmethod
    def _add_heading_safe(document, text, level=1):
        style_by_level = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
        style_name = style_by_level.get(level)
        if style_name:
            try:
                return document.add_paragraph(text, style=style_name)
            except Exception:
                pass
        paragraph = document.add_paragraph(text)
        if level == 1:
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.size = None
        elif level == 2:
            paragraph.runs[0].bold = True
        return paragraph

    @staticmethod
    def _decode_base64_image(content_base64):
        payload = content_base64.split(',', 1)[1] if ',' in content_base64 else content_base64
        try:
            return base64.b64decode(payload, validate=True)
        except (ValueError, binascii.Error) as exc:
            raise RuntimeError("No se pudo decodificar un fotograma en base64.") from exc

    @staticmethod
    def _prepare_picture_stream(frame):
        raw = IntegrityService._decode_base64_image(frame.get('content_base64', ''))
        mime = frame.get('mime_type')
        if mime != 'image/webp':
            return BytesIO(raw)

        try:
            from PIL import Image
        except ImportError as exc:
            raise RuntimeError("Falta Pillow para procesar imagenes WEBP.") from exc

        try:
            source = BytesIO(raw)
            target = BytesIO()
            with Image.open(source) as image:
                image.convert('RGB').save(target, format='PNG')
            target.seek(0)
            return target
        except Exception as exc:
            raise RuntimeError("No se pudo convertir un fotograma WEBP.") from exc

    @staticmethod
    def _append_frames_annex(document, frames):
        if not frames:
            return

        try:
            from docx.shared import Inches
        except ImportError as exc:
            raise RuntimeError("python-docx no permite insertar imagenes (docx.shared no disponible).") from exc

        ordered = sorted(frames, key=lambda frame: frame.get('order', 0))
        document.add_page_break()
        IntegrityService._add_heading_safe(document, "Anexo de fotogramas", level=1)

        for idx, frame in enumerate(ordered, start=1):
            file_name = frame.get('file_name', f'fotograma_{idx}.jpg')
            description = frame.get('description') or ''
            IntegrityService._add_heading_safe(document, f"Fotograma {idx}: {file_name}", level=2)
            pic_stream = IntegrityService._prepare_picture_stream(frame)
            try:
                document.add_picture(pic_stream, width=Inches(6.2))
            except Exception as exc:
                raise RuntimeError(f"No se pudo insertar el fotograma '{file_name}'.") from exc

            if description:
                document.add_paragraph(f"Descripcion: {description}")

            document.add_paragraph("")

    @staticmethod
    def _normalize_video_payload(payload):
        # Nuevo formato recomendado: {'report_data': {...}, 'frames': [...]}
        if 'report_data' in payload:
            return payload.get('report_data') or {}, payload.get('frames') or []
        # Compatibilidad legacy
        return payload, []

    @staticmethod
    def _replace_text_in_docx(document, replacements):
        def replace_in_paragraph(paragraph):
            original = paragraph.text
            updated = original
            for src, dst in replacements:
                updated = updated.replace(src, dst)
            if updated != original:
                paragraph.text = updated

        for paragraph in document.paragraphs:
            replace_in_paragraph(paragraph)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                replace_in_paragraph(paragraph)
            for paragraph in section.footer.paragraphs:
                replace_in_paragraph(paragraph)

    @staticmethod
    def _extract_json_object(raw_text):
        if isinstance(raw_text, dict):
            return raw_text

        text = str(raw_text or '').strip()
        if not text:
            raise RuntimeError("La API de IA no devolvio contenido para mejorar el texto.")

        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

        match = re.search(r"\{.*\}", text, re.DOTALL)
        if not match:
            raise RuntimeError("No se pudo interpretar la respuesta de la API de IA.")

        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError as exc:
            raise RuntimeError("No se pudo interpretar la respuesta JSON de la API de IA.") from exc

    @staticmethod
    def improve_report_text_with_ai(desarrollo, conclusion, custom_api_key=""):
        api_key_from_settings = str(getattr(settings, 'AI_TEXT_API_KEY', '') or '').strip()
        api_key = custom_api_key.strip() or api_key_from_settings
        api_url = str(getattr(settings, 'AI_TEXT_API_URL', '') or '').strip()
        model = str(getattr(settings, 'AI_TEXT_MODEL', '') or '').strip()
        timeout_seconds = int(getattr(settings, 'AI_TEXT_TIMEOUT_SECONDS', 45))

        if not api_key:
            raise RuntimeError("No se configuro AI_TEXT_API_KEY y no se ingreso ninguna manualmente.")
        if not api_url:
            raise RuntimeError("No se configuro AI_TEXT_API_URL.")
        if not model:
            raise RuntimeError("No se configuro AI_TEXT_MODEL.")

        system_prompt = (
            "Eres un redactor profesional de informes formales en espanol. "
            "Mejora redaccion, coherencia, ortografia y puntuacion, sin inventar hechos nuevos. "
            "Mantener un tono tecnico y objetivo."
        )

        payload_prompt = {
            "instrucciones": "Mejora el texto y devuelve SOLO JSON valido.",
            "formato_estricto": {
                "desarrollo": "texto mejorado",
                "conclusion": "texto mejorado",
            },
            "desarrollo_original": str(desarrollo or ''),
            "conclusion_original": str(conclusion or ''),
        }

        request_payload = {
            "model": model,
            "temperature": 0.2,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": json.dumps(payload_prompt, ensure_ascii=False)},
            ],
            "response_format": {"type": "json_object"},
        }

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }

        try:
            response = requests.post(
                api_url,
                headers=headers,
                json=request_payload,
                timeout=timeout_seconds,
            )
        except requests.RequestException as exc:
            raise RuntimeError("No se pudo conectar con la API de IA.") from exc

        if response.status_code >= 400:
            detail = ""
            try:
                error_payload = response.json()
                error_node = error_payload.get('error', {})
                if isinstance(error_node, dict):
                    detail = str(error_node.get('message', '') or '')
                elif error_node:
                    detail = str(error_node)
            except Exception:
                detail = response.text[:200]

            suffix = f" Detalle: {detail}" if detail else ""
            raise RuntimeError(
                f"La API de IA devolvio un error ({response.status_code}).{suffix}"
            )

        try:
            response_payload = response.json()
            choices = response_payload.get('choices') or []
            if not choices:
                raise RuntimeError("La API de IA no devolvio opciones de respuesta.")

            message = choices[0].get('message') if isinstance(choices[0], dict) else None
            content = message.get('content') if isinstance(message, dict) else None
            if isinstance(content, list):
                parts = []
                for item in content:
                    if isinstance(item, dict):
                        parts.append(str(item.get('text', '') or ''))
                    else:
                        parts.append(str(item))
                content = ''.join(parts)

            parsed_content = IntegrityService._extract_json_object(content)
            improved_desarrollo = IntegrityService._as_text(
                parsed_content.get('desarrollo'),
                str(desarrollo or '')
            ).strip()
            improved_conclusion = IntegrityService._as_text(
                parsed_content.get('conclusion'),
                str(conclusion or '')
            ).strip()

            return {
                'desarrollo': improved_desarrollo,
                'conclusion': improved_conclusion,
            }
        except RuntimeError:
            raise
        except Exception as exc:
            raise RuntimeError("Respuesta invalida al mejorar texto con IA.") from exc

    @staticmethod
    def generate_video_analysis_docx(payload):
        """
        Genera un informe DOCX usando data/Modelo Informe.docx como base.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx no esta instalado") from exc

        template_path = Path(__file__).resolve().parents[3] / "data" / "Modelo Informe.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"No se encontro plantilla: {template_path}")

        report_data, frames = IntegrityService._normalize_video_payload(payload or {})

        report_date = report_data.get("report_date") or datetime.now().strftime("%Y-%m-%d")
        try:
            if isinstance(report_date, datetime):
                parsed = report_date
            else:
                parsed = datetime.strptime(str(report_date), "%Y-%m-%d")
        except ValueError:
            parsed = datetime.now()

        month_names = {
            1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
            5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
            9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
        }
        month_name = month_names.get(parsed.month, "enero")
        year = str(parsed.year)

        destinatarios = IntegrityService._as_text(report_data.get("destinatarios"), "URSA I - Jefe")
        tipo_informe = IntegrityService._as_text(report_data.get("tipo_informe"), "IAV - Informe Analisis de Video")
        numero_informe = IntegrityService._as_text(report_data.get("numero_informe"), f"001/{parsed.year}")
        grado = IntegrityService._as_text(report_data.get("grado"), "Oficial Mayor")
        operador = IntegrityService._as_text(report_data.get("operador"), "Operador")
        lup = IntegrityService._as_text(report_data.get("lup"), "0000")
        sistema = IntegrityService._as_text(report_data.get("sistema"), "MILESTONE")
        prevencion_sumaria = IntegrityService._as_text(report_data.get("prevencion_sumaria"), "000BAR/2026")
        caratula = IntegrityService._as_text(report_data.get("caratula"), "DENUNCIA S/ PRESUNTO HURTO")
        unidad = IntegrityService._as_text(report_data.get("unidad"), "Coordinación Regional de Video Seguridad I del Este")
        material_filmico = IntegrityService._as_text(report_data.get("material_filmico"), "material fílmico")
        fiscalia = IntegrityService._as_text(report_data.get("fiscalia"), "Fiscalia Nro. 02")
        fiscal = IntegrityService._as_text(report_data.get("fiscal"), "Fiscal Interviniente")
        denunciante = IntegrityService._as_text(report_data.get("denunciante"), "Denunciante")
        vuelo = IntegrityService._as_text(report_data.get("vuelo"), "---")
        empresa_aerea = IntegrityService._as_text(report_data.get("empresa_aerea"), "---")
        destino = IntegrityService._as_text(report_data.get("destino"), "---")
        fecha_hecho = IntegrityService._as_text(report_data.get("fecha_hecho"), "---")
        unidad_aeroportuaria = IntegrityService._as_text(report_data.get("unidad_aeroportuaria"), "---")
        asiento = IntegrityService._as_text(report_data.get("asiento"), "---")
        aeropuerto = IntegrityService._as_text(report_data.get("aeropuerto"), "---")
        objeto_denunciado = IntegrityService._as_text(report_data.get("objeto_denunciado"), "objeto denunciado")
        conclusion = IntegrityService._as_text(report_data.get("conclusion"), "")
        desarrollo = IntegrityService._as_text(report_data.get("desarrollo"), "")
        firma = IntegrityService._as_text(report_data.get("firma"), "Coordinador CReV I DEL ESTE")

        op_info = f"{grado} {operador}, LUP: {lup}"

        replacements = [
            ("Fecha: Enero de 2026.", f"Fecha: {month_name.capitalize()} de {year}."),
            ("URSA I - Jefe", destinatarios),
            ("IAV - Informe Análisis de Video    /2026", f"{tipo_informe} {numero_informe}"),
            ("IAV - Informe Analisis de Video    /2026", f"{tipo_informe} {numero_informe}"),
            ("Oficial Mayor XXX, LUP: XXX", op_info),
            ("denominado “MILESTONE”", f"denominado \"{sistema}\""),
            ("denominado \"MILESTONE\"", f"denominado \"{sistema}\""),
            ("Prevención Sumaria 003BAR/2026", f"Prevencion Sumaria {prevencion_sumaria}"),
            ("“DENUNCIA S/ PRESUNTO HURTO”", f"\"{caratula}\""),
            ("\"DENUNCIA S/ PRESUNTO HURTO\"", f"\"{caratula}\""),
            ("Fiscalía Nro. 02 de San Carlos de Bariloche, Provincia de Rio Negro", fiscalia),
            ("Dr. INTI ISLA", f"Dr. {fiscal}"),
            ("Sr. PONZO Osvaldo", f"Sr. {denunciante}"),
            ("vuelo WJ 3045", f"vuelo {vuelo}"),
            ("RIVER PLATE", objeto_denunciado),
            ("Coordinación Regional de Video Seguridad I del Este", unidad),
            ("material fílmico", material_filmico),
            ("Coordinador CReV I DEL ESTE", firma),
            ("empresa aerea", empresa_aerea),
            ("empresa aérea", empresa_aerea),
            ("destino", destino),
            ("fecha del hecho", fecha_hecho),
            ("unidad aeroportuaria", unidad_aeroportuaria),
            ("asiento", asiento),
            ("aeropuerto", aeropuerto),
        ]

        document = Document(str(template_path))
        IntegrityService._replace_text_in_docx(document, replacements)

        if desarrollo or conclusion:
            document.add_page_break()
            IntegrityService._add_heading_safe(document, "Ampliacion del informe", level=1)
            if desarrollo:
                IntegrityService._add_heading_safe(document, "Desarrollo", level=2)
                document.add_paragraph(desarrollo)
            if conclusion:
                IntegrityService._add_heading_safe(document, "Conclusion", level=2)
                document.add_paragraph(conclusion)

        IntegrityService._append_frames_annex(document, frames)

        out = BytesIO()
        document.save(out)
        out.seek(0)
        return out, f"informe_analisis_video_{parsed.strftime('%Y%m%d')}.docx"

    @staticmethod
    def _get_hasher(algorithm='sha256'):
        if algorithm == 'sha256':
            return hashlib.sha256()
        if algorithm == 'sha512':
            return hashlib.sha512()
        if algorithm == 'sha3':
            return hashlib.sha3_256()
        raise ValueError(f"Algoritmo no soportado: {algorithm}")

    @staticmethod
    def calculate_file_hash(file_obj, algorithm='sha256'):
        """
        Calcula el hash de un archivo usando el algoritmo especificado.
        
        Args:
            file_obj: Objeto de archivo (Django UploadedFile o file-like object)
            algorithm: 'sha256', 'sha512' o 'sha3' (default: sha256)
        
        Returns:
            str: Hash hexadecimal del archivo
        """
        hasher = IntegrityService._get_hasher(algorithm)

        # Leer archivo en chunks para evitar problemas de memoria
        for chunk in file_obj.chunks():
            hasher.update(chunk)
            
        return hasher.hexdigest()
    
    @staticmethod
    def verify_existing_backup(backup_path, expected_hash, algorithm='sha256'):
        """
        Verifica la integridad de un archivo de backup existente.
        
        Args:
            backup_path: Ruta completa al archivo de backup
            expected_hash: Hash esperado para comparación
            algorithm: Algoritmo de hash utilizado
        
        Returns:
            dict: {'is_valid': bool, 'calculated_hash': str, 'message': str}
        """
        if not os.path.exists(backup_path):
            return {
                'is_valid': False,
                'calculated_hash': None,
                'message': f'El archivo no existe en la ruta: {backup_path}'
            }
        
        try:
            hasher = IntegrityService._get_hasher(algorithm)
            
            # Leer archivo en chunks
            with open(backup_path, 'rb') as f:
                while chunk := f.read(8192):
                    hasher.update(chunk)
            
            calculated_hash = hasher.hexdigest()
            is_valid = calculated_hash == expected_hash
            
            return {
                'is_valid': is_valid,
                'calculated_hash': calculated_hash,
                'message': 'Integridad verificada correctamente' if is_valid else 'El hash no coincide - archivo modificado'
            }
        except Exception as e:
            return {
                'is_valid': False,
                'calculated_hash': None,
                'message': f'Error al verificar archivo: {str(e)}'
            }

    @staticmethod
    def generate_integrity_pdf(file_name, file_size, hashes):
        """
        Genera un reporte PDF básico con detalles del archivo y hashes calculados.
        
        Args:
            file_name: Nombre del archivo
            file_size: Tamaño en bytes
            hashes: dict {'SHA256': '...', 'SHA512': '...', 'SHA-3': '...'}
        
        Returns:
            BytesIO: Buffer con el PDF generado
        """
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # Header
        p.setFont("Helvetica-Bold", 18)
        p.drawString(50, height - 50, "Reporte de Integridad de Archivo")
        
        p.setFont("Helvetica", 10)
        p.drawString(50, height - 70, f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        p.line(50, height - 80, width - 50, height - 80)

        # File Details
        y = height - 120
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, "Detalles del Archivo")
        y -= 20
        p.setFont("Helvetica", 12)
        p.drawString(70, y, f"Nombre: {file_name}")
        y -= 20
        p.drawString(70, y, f"Tamaño: {file_size} bytes")
        
        y -= 40
        p.line(50, y, width - 50, y)
        y -= 30

        # Hashes
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, "Verificación de Integridad (Hashes)")
        y -= 30

        for algo, hash_val in hashes.items():
            p.setFont("Helvetica-Bold", 11)
            p.drawString(70, y, f"{algo}:")
            p.setFont("Courier", 10)
            p.drawString(130, y, hash_val)
            y -= 25

        # Footer
        p.setFont("Helvetica-Oblique", 8)
        p.drawString(50, 50, "Este documento es un reporte generado automáticamente por el sistema GestorCOC.")
        p.drawString(50, 40, "No garantiza la autenticidad del archivo si este es modificado posteriormente.")

        p.showPage()
        p.save()
        
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_integrity_summary_pdf(entries):
        """
        Genera un PDF resumen en orientación landscape con todos los archivos
        hasheados en la sesión UI. El hash se muestra completo.

        Args:
            entries: list[dict] con name, size, algorithm, hash, time_ms

        Returns:
            BytesIO: Buffer con el PDF generado
        """
        buffer = BytesIO()
        page_size = landscape(A4)
        p = canvas.Canvas(buffer, pagesize=page_size)
        width, height = page_size

        # Márgenes y ancho útil
        margin = 40
        usable_width = width - (margin * 2)

        def draw_header():
            p.setFont("Helvetica-Bold", 16)
            p.drawString(margin, height - 35, "Resumen de Verificación de Integridad")
            p.setFont("Helvetica", 9)
            p.drawString(margin, height - 52, f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}   |   Total archivos: {len(entries)}")
            p.line(margin, height - 60, width - margin, height - 60)

        def draw_column_headers(y_pos):
            p.setFillColorRGB(0.11, 0.22, 0.37)  # Azul oscuro
            p.rect(margin, y_pos - 4, usable_width, 18, fill=1, stroke=0)
            p.setFillColorRGB(1, 1, 1)
            p.setFont("Helvetica-Bold", 8)
            p.drawString(margin + 4,      y_pos + 2, "#")
            p.drawString(margin + 20,     y_pos + 2, "Nombre")
            p.drawString(margin + 170,    y_pos + 2, "Tipo")
            p.drawString(margin + 210,    y_pos + 2, "Tamaño")
            p.drawString(margin + 260,    y_pos + 2, "Algoritmo")
            p.drawString(margin + 310,    y_pos + 2, "Tiempo")
            p.drawString(margin + 360,    y_pos + 2, "Hash completo")
            p.setFillColorRGB(0, 0, 0)
            return y_pos - 20

        draw_header()
        y = height - 80
        y = draw_column_headers(y)

        # Ancho máximo del campo hash en puntos (desde x=400 hasta el margen derecho)
        hash_x = margin + 360
        hash_max_width = width - margin - hash_x
        # SHA-256 = 64 chars, SHA-512 = 128 chars, SHA-3 = 64 chars
        # Con Courier 7, un car es ~4.2pt → caben ~hash_max_width/4.2 chars por línea
        chars_per_line = int(hash_max_width / 4.2)

        row_bg_even = (0.97, 0.98, 1.0)
        row_bg_odd  = (1.0, 1.0, 1.0)

        for idx, entry in enumerate(entries, start=1):
            name       = str(entry.get("name", "N/A"))
            ext_type   = str(entry.get("type", "N/A"))
            size_str   = str(entry.get("size_str", "N/A"))
            algorithm  = str(entry.get("algorithm", "N/A"))
            time_ms    = str(entry.get("time_ms", "N/A")).strip()
            hash_value = str(entry.get("hash", "N/A"))

            # Quitar posibles dobles ms si la UI los envía por error
            if time_ms.endswith(" ms ms"):
                time_ms = time_ms[:-3]

            # Partir el hash en líneas si es muy largo
            hash_lines = [hash_value[i:i + chars_per_line] for i in range(0, len(hash_value), chars_per_line)] if hash_value != "N/A" else [hash_value]
            row_height = max(14, len(hash_lines) * 10 + 4)

            if y - row_height < margin + 20:
                p.showPage()
                draw_header()
                y = height - 80
                y = draw_column_headers(y)

            # Fondo alternado
            bg = row_bg_even if idx % 2 == 0 else row_bg_odd
            p.setFillColorRGB(*bg)
            p.rect(margin, y - row_height + 12, usable_width, row_height, fill=1, stroke=0)
            p.setFillColorRGB(0, 0, 0)

            # Datos de la fila
            row_top = y
            p.setFont("Helvetica", 8)
            p.drawString(margin + 4,   row_top, str(idx))
            p.drawString(margin + 20,  row_top, name[:25] if len(name) > 25 else name)
            p.drawString(margin + 170, row_top, ext_type[:7] if len(ext_type) > 7 else ext_type)
            p.drawString(margin + 210, row_top, size_str)
            p.drawString(margin + 260, row_top, algorithm)
            p.drawString(margin + 310, row_top, str(time_ms))

            # Hash completo en líneas
            p.setFont("Courier", 7)
            p.setFillColorRGB(0.21, 0.19, 0.64)  # Índigo
            for line_idx, hash_line in enumerate(hash_lines):
                p.drawString(hash_x, row_top - (line_idx * 9), hash_line)
            p.setFillColorRGB(0, 0, 0)

            # Línea separadora
            p.setStrokeColorRGB(0.88, 0.91, 0.94)
            p.line(margin, y - row_height + 12, width - margin, y - row_height + 12)
            p.setStrokeColorRGB(0, 0, 0)

            y -= row_height

        # Footer
        p.setFont("Helvetica-Oblique", 7)
        p.setFillColorRGB(0.58, 0.64, 0.72)
        p.drawString(margin, margin - 10, "Documento generado automáticamente por GestorCOC. Los hashes son calculados por el cliente y no garantizan la cadena de custodia si los archivos son modificados posteriormente.")
        p.setFillColorRGB(0, 0, 0)

        p.showPage()
        p.save()
        buffer.seek(0)
        return buffer

    
    @staticmethod
    def generate_verification_report(film_record):
        """
        Genera un reporte completo de verificación para un FilmRecord.
        Incluye todos los datos del registro, verificación CREV y hash de integridad.
        
        Args:
            film_record: Instancia de FilmRecord
        
        Returns:
            BytesIO: Buffer con el PDF de certificación
        """
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # Header
        p.setFont("Helvetica-Bold", 20)
        p.drawString(50, height - 50, "CERTIFICADO DE INTEGRIDAD")
        p.drawString(50, height - 75, "Registro Fílmico - GestorCOC")
        
        p.setFont("Helvetica", 9)
        p.drawString(50, height - 95, f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        p.line(50, height - 105, width - 50, height - 105)

        # Información del Registro
        y = height - 140
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, y, f"Registro N° {film_record.id}")
        y -= 30
        
        # Datos principales
        p.setFont("Helvetica-Bold", 11)
        p.drawString(50, y, "Información de la Solicitud:")
        y -= 20
        p.setFont("Helvetica", 10)
        
        fields = [
            ("Nº Asunto", film_record.issue_number),
            ("Nº Causa Judicial", film_record.judicial_case_number),
            ("Solicitante", film_record.requester),
            ("Fecha de Ingreso", film_record.entry_date.strftime('%d/%m/%Y') if film_record.entry_date else 'N/A'),
            ("Cámara", film_record.camera.name if film_record.camera else 'N/A'),
            ("Período", f"{film_record.start_time.strftime('%d/%m/%Y %H:%M')} - {film_record.end_time.strftime('%d/%m/%Y %H:%M')}"),
        ]
        
        for label, value in fields:
            if value:
                p.drawString(70, y, f"{label}: {value}")
                y -= 18
        
        # Información de Backup
        y -= 20
        p.setFont("Helvetica-Bold", 11)
        p.drawString(50, y, "Información de Backup:")
        y -= 20
        p.setFont("Helvetica", 10)
        p.drawString(70, y, f"Ruta: {film_record.backup_path or 'No especificada'}")
        y -= 18
        p.drawString(70, y, f"Tamaño: {film_record.file_size or 'N/A'} bytes")
        
        # Hash de Integridad
        y -= 30
        p.setFont("Helvetica-Bold", 11)
        p.drawString(50, y, "Hash de Integridad (SHA-256):")
        y -= 20
        p.setFont("Courier", 8)
        p.drawString(70, y, film_record.file_hash or 'No calculado')
        
        # Verificación CREV
        y -= 40
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, "VERIFICACIÓN CREV")
        p.line(50, y - 5, width - 50, y - 5)
        y -= 25
        p.setFont("Helvetica", 10)
        
        if film_record.verified_by_crev:
            p.drawString(70, y, f"Verificado por: {film_record.verified_by_crev.last_name}, {film_record.verified_by_crev.first_name}")
            y -= 18
            p.drawString(70, y, f"Fecha de Verificación: {film_record.verification_date.strftime('%d/%m/%Y %H:%M:%S')}")
            y -= 18
            p.setFont("Helvetica-Bold", 10)
            p.setFillColorRGB(0, 0.5, 0)  # Verde
            p.drawString(70, y, "✓ REGISTRO CERTIFICADO Y BLOQUEADO PARA EDICIÓN")
        else:
            p.setFillColorRGB(0.8, 0, 0)  # Rojo
            p.drawString(70, y, "✗ PENDIENTE DE VERIFICACIÓN CREV")
        
        # Footer
        p.setFillColorRGB(0, 0, 0)  # Negro
        p.setFont("Helvetica-Oblique", 8)
        p.drawString(50, 50, "Este certificado es válido solo si el hash coincide con el archivo de backup.")
        p.drawString(50, 40, "Sistema GestorCOC - Centro de Operaciones y Control")

        p.showPage()
        p.save()
        
        buffer.seek(0)
        return buffer
